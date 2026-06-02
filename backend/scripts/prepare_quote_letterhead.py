#!/usr/bin/env python3
"""Prepare a quote letterhead .docx with indexed merge codes for fee lines.

Usage:
  backend/.venv/bin/python backend/scripts/prepare_quote_letterhead.py \\
    --src ~/Downloads/Letterhead.docx \\
    --dest ~/Downloads/Letterhead.docx
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _clear_body_keep_section(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def prepare_quote_letterhead(src: Path, dest: Path, *, slots: int = 25) -> None:
    doc = Document(str(src))
    _clear_body_keep_section(doc)

    doc.add_paragraph("[ORG_AND_ADDRESS_BLOCK]")
    doc.add_paragraph("")
    doc.add_paragraph(f"[DATE]")
    doc.add_paragraph("")
    doc.add_paragraph("Re: [MATTER_DESCRIPTION] — [CASE_REF]")
    doc.add_paragraph("")
    doc.add_paragraph("Dear [CONTACT_LETTER_DEAR]")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Thank you for instructing us. Set out below is our estimate of costs based on a property value of "
        "[QUOTE_PROPERTY_VALUE]."
    )
    doc.add_paragraph("")

    table = doc.add_table(rows=1 + slots, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Description"
    hdr[1].text = "Amount"
    hdr[2].text = "VAT"
    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    for i in range(1, slots + 1):
        tag = f"{i:02d}"
        row = table.rows[i].cells
        row[0].text = f"[QUOTE_{tag}_LABEL]"
        amt = row[1].paragraphs[0]
        amt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        amt.text = f"[QUOTE_{tag}_AMOUNT]"
        vat_cell = row[2].paragraphs[0]
        vat_cell.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        vat_cell.text = f"[QUOTE_{tag}_VAT]"

    doc.add_paragraph("")
    closing = doc.add_paragraph("Yours faithfully")
    closing.runs[0].font.size = Pt(11)
    doc.add_paragraph("")
    doc.add_paragraph("[FEE_EARNER]")
    doc.add_paragraph("[FIRM_TRADING_NAME]")

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest))


def main() -> None:
    parser = argparse.ArgumentParser(description="Prepare quote letterhead with merge codes")
    parser.add_argument("--src", type=Path, required=True)
    parser.add_argument("--dest", type=Path, required=True)
    parser.add_argument("--slots", type=int, default=25)
    args = parser.parse_args()
    prepare_quote_letterhead(args.src, args.dest, slots=args.slots)
    print(f"Wrote {args.dest}")


if __name__ == "__main__":
    main()
