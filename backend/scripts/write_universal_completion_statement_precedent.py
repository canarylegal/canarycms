#!/usr/bin/env python3
"""Emit the canonical universal completion statement template .docx (reserved COMPLETION_STATEMENT precedent).

Run manually; the bundled seed copy lives at
``backend/precedents_seed/bundle/g2_completion_statement.docx``.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def write_universal_completion_statement_precedent(path: Path, *, slots: int = 50) -> None:
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("COMPLETION STATEMENT")
    run.bold = True
    run.font.size = Pt(16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("[CASE_REF] — [MATTER_DESCRIPTION]").font.size = Pt(11)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run("Date: [COMPLETION_DATE]").font.size = Pt(10)

    doc.add_paragraph("")

    table = doc.add_table(rows=1 + slots + 2, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, label in enumerate(("Description", "Debit", "Credit")):
        hdr[i].text = label
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
            if i > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for i in range(1, slots + 1):
        tag = f"{i:02d}"
        row = table.rows[i].cells
        row[0].text = f"[COMPLETION_{tag}_DESCRIPTION]"
        row[1].text = f"[COMPLETION_{tag}_DEBIT]"
        row[2].text = f"[COMPLETION_{tag}_CREDIT]"
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    totals_row = table.rows[slots + 1].cells
    totals_row[0].text = "TOTALS"
    totals_row[1].text = "[COMPLETION_TOTAL_DEBIT]"
    totals_row[2].text = "[COMPLETION_TOTAL_CREDIT]"
    for cell in totals_row:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    totals_row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    totals_row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    balance_row = table.rows[slots + 2].cells
    balance_row[0].merge(balance_row[1])
    balance_row[0].text = "[COMPLETION_BALANCE_LABEL]"
    balance_row[2].text = "[COMPLETION_BALANCE_AMOUNT]"
    for cell in balance_row:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    balance_row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def main() -> None:
    out = (
        Path(sys.argv[1]).expanduser().resolve()
        if len(sys.argv) > 1
        else Path("precedents_seed/bundle/g2_completion_statement.docx")
    )
    write_universal_completion_statement_precedent(out)
    print(out)


if __name__ == "__main__":
    main()
