#!/usr/bin/env python3
"""Emit the canonical universal invoice template .docx (reserved INVOICE_TEMPLATE precedent).

Run manually; the bundled seed copy lives at
``backend/precedents_seed/bundle/g1_invoice_template.docx``.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def write_universal_invoice_precedent(path: Path, *, slots: int = 25) -> None:
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("INVOICE")
    run.bold = True
    run.font.size = Pt(16)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run("[INVOICE_NUMBER]  ·  [INVOICE_DATE]")
    meta_run.font.size = Pt(11)

    doc.add_paragraph("")

    bill = doc.add_paragraph()
    bill.add_run("Bill to: ").bold = True
    bill.add_run("[INVOICE_BILL_TO]")

    doc.add_paragraph("Matter: [CASE_REF] — [MATTER_DESCRIPTION]")
    doc.add_paragraph("Fee earner: [FEE_EARNER]")
    doc.add_paragraph("")

    table = doc.add_table(rows=1 + slots, cols=5)
    table.style = "Table Grid"
    headers = ("Type", "Description", "Net", "VAT", "Total")
    hdr = table.rows[0].cells
    for i, label in enumerate(headers):
        hdr[i].text = label
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
            if i >= 2:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for i in range(1, slots + 1):
        tag = f"{i:02d}"
        row = table.rows[i].cells
        row[0].text = f"[INVOICE_{tag}_TYPE]"
        row[1].text = f"[INVOICE_{tag}_DESCRIPTION]"
        row[2].text = f"[INVOICE_{tag}_NET]"
        row[3].text = f"[INVOICE_{tag}_VAT]"
        row[4].text = f"[INVOICE_{tag}_TOTAL]"
        for j in range(2, 5):
            row[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph("")

    totals = doc.add_table(rows=3, cols=2)
    totals.style = "Table Grid"
    for i, (label, code) in enumerate(
        (
            ("Net total", "[INVOICE_NET_TOTAL]"),
            ("VAT total", "[INVOICE_VAT_TOTAL]"),
            ("Invoice total", "[INVOICE_TOTAL]"),
        )
    ):
        totals.rows[i].cells[0].text = label
        amt = totals.rows[i].cells[1]
        amt.text = code
        amt.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if i == 2:
            for cell in totals.rows[i].cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def main() -> None:
    out = (
        Path(sys.argv[1]).expanduser().resolve()
        if len(sys.argv) > 1
        else Path("precedents_seed/bundle/g1_invoice_template.docx")
    )
    write_universal_invoice_precedent(out)
    print(out)


if __name__ == "__main__":
    main()
