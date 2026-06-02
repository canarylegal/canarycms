#!/usr/bin/env python3
"""Add a VAT column and total merge codes to an existing quote letterhead .docx.

Preserves headers, footers, and body content (unlike prepare_quote_letterhead.py which rebuilds the body).

Usage:
  backend/.venv/bin/python backend/scripts/upgrade_quote_letterhead_vat_column.py \\
    --path ~/Downloads/Letterhead.docx
"""

from __future__ import annotations

import argparse
import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

_QUOTE_SLOT_RE = re.compile(r"\[QUOTE_(\d{2})_(?:LABEL|AMOUNT)\]", re.I)
_DEFAULT_SLOTS = 25


def _find_quote_table(doc: Document):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if _QUOTE_SLOT_RE.search(cell.text or ""):
                    return table
    return doc.tables[0] if doc.tables else None


def _add_vat_column(table, *, slots: int = _DEFAULT_SLOTS) -> None:
    if len(table.columns) >= 3:
        return
    tbl = table._tbl
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        grid.append(OxmlElement("w:gridCol"))

    for row in table.rows:
        tr = row._tr
        ref_tc = row.cells[1]._tc
        tc = OxmlElement("w:tc")
        tcPr = ref_tc.find(qn("w:tcPr"))
        if tcPr is not None:
            tc.append(deepcopy(tcPr))
        tc.append(OxmlElement("w:p"))
        tr.append(tc)

    hdr = table.rows[0].cells
    hdr[2].text = "VAT"
    for p in hdr[2].paragraphs:
        for run in p.runs:
            run.bold = True

    for i in range(1, min(slots + 1, len(table.rows))):
        tag = f"{i:02d}"
        vat_para = table.rows[i].cells[2].paragraphs[0]
        vat_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        vat_para.text = f"[QUOTE_{tag}_VAT]"


def _ensure_total_merge_paragraphs(doc: Document) -> None:
    if any("[QUOTE_GRAND_TOTAL]" in (p.text or "") for p in doc.paragraphs):
        return
    target = None
    for p in doc.paragraphs:
        if (p.text or "").strip().startswith("Yours faithfully"):
            target = p
            break
    if target is None:
        return
    parent = target._element.getparent()
    idx = list(parent).index(target._element)
    for offset, line in enumerate(
        [
            "",
            "Total (main column): [QUOTE_MAIN_TOTAL]",
            "Total VAT: [QUOTE_VAT_TOTAL]",
            "Total estimated costs: [QUOTE_GRAND_TOTAL]",
            "",
        ]
    ):
        new_p = OxmlElement("w:p")
        if line:
            r = OxmlElement("w:r")
            t = OxmlElement("w:t")
            t.text = line
            r.append(t)
            new_p.append(r)
        parent.insert(idx + offset, new_p)


def upgrade_quote_letterhead(path: Path, *, slots: int = _DEFAULT_SLOTS) -> None:
    doc = Document(str(path))
    table = _find_quote_table(doc)
    if table is None:
        raise ValueError("No table found in letterhead.")
    _add_vat_column(table, slots=slots)
    _ensure_total_merge_paragraphs(doc)
    doc.save(str(path))


def main() -> None:
    parser = argparse.ArgumentParser(description="Add VAT column to quote letterhead")
    parser.add_argument("--path", type=Path, required=True, help="Letterhead .docx to update in place")
    parser.add_argument("--slots", type=int, default=_DEFAULT_SLOTS)
    args = parser.parse_args()
    upgrade_quote_letterhead(args.path.resolve(), slots=args.slots)
    print(f"Updated {args.path}")


if __name__ == "__main__":
    main()
