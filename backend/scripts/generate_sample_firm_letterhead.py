"""Emit a fictitious firm letterhead .docx (header + footer) for Admin → Firm → digital letterhead testing.

Run from repo root or backend:

  python scripts/generate_sample_firm_letterhead.py
  python scripts/generate_sample_firm_letterhead.py /path/out.docx

Canary copies header/footer XML from this file onto composed letters (see ``apply_digital_letterhead_headers_footers``).
Embedded header/footer logos are supported.
"""

from __future__ import annotations

import io
import sys
from pathlib import Path

_ROOT = Path(__file__).resolve().parents[1]
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.docx_util import ensure_docx_proofing_language_en_gb_bytes


def _border_bottom_paragraph(para, *, sz: str = "12", color: str = "2F5496") -> None:
    p_pr = para._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_header_footer(*, default_out: Path) -> Path:
    out = (
        Path(sys.argv[1]).expanduser().resolve()
        if len(sys.argv) > 1
        else default_out
    )

    doc = Document()
    # Single empty body paragraph — upload target is header/footer only.
    doc.add_paragraph("")

    section = doc.sections[0]
    # Section geometry must leave enough room *between* header/footer offsets and body margins.
    # OOXML reserves header vertical space ≈ top_margin − header_distance; too small a gap
    # jams multi-line headers against the body (Word may reflow; ONLYOFFICE often does not).
    section.header_distance = Inches(0.55)
    section.footer_distance = Inches(0.55)
    section.top_margin = Inches(1.65)
    section.bottom_margin = Inches(1.15)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    hdr = section.header
    hdr.paragraphs[0].clear()

    firm_block = [
        ("BLACKTHORN ELLIS SOLICITORS LLP", True, Pt(13), RGBColor(0x2F, 0x54, 0x96)),
        ("Commercial · Property · Private Client", False, Pt(9), RGBColor(0x55, 0x55, 0x55)),
        ("47 Charterhouse Yard, London EC1M 6EA", False, Pt(9), RGBColor(0x33, 0x33, 0x33)),
        ("DX: DX 99999 MOORGATE", False, Pt(9), RGBColor(0x33, 0x33, 0x33)),
        (
            "T +44 (0)20 7946 0123   ·   F +44 (0)20 7946 0124   ·   enquiries@blackthornellis.example",
            False,
            Pt(9),
            RGBColor(0x33, 0x33, 0x33),
        ),
    ]
    for i, (text, bold, size, colour) in enumerate(firm_block):
        para = hdr.paragraphs[0] if i == 0 else hdr.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(text)
        run.bold = bold
        run.font.size = size
        run.font.color.rgb = colour
        run.font.name = "Calibri"

    _border_bottom_paragraph(hdr.paragraphs[-1])

    ftr = section.footer
    ftr.paragraphs[0].clear()

    footer_lines: list[tuple[str, bool]] = [
        (
            "Blackthorn Ellis Solicitors LLP is a limited liability partnership registered in England and Wales "
            "(OC999999).",
            False,
        ),
        (
            "Registered office: 47 Charterhouse Yard, London EC1M 6EA. Professional indemnity insurance is maintained "
            "in accordance with SRA rules; further details on request.",
            False,
        ),
        (
            "Authorised and regulated by the Solicitors Regulation Authority | "
            "SRA ID 999999 | VAT GB 123 4567 89",
            False,
        ),
        (
            "www.blackthornellis.example — entirely fictitious letterhead for software testing",
            True,
        ),
    ]
    for i, (text, bold) in enumerate(footer_lines):
        para = ftr.paragraphs[0] if i == 0 else ftr.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.font.name = "Calibri"
        run.bold = bold

    buf = io.BytesIO()
    doc.save(buf)
    patched = ensure_docx_proofing_language_en_gb_bytes(buf.getvalue())

    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_bytes(patched)
    return out


def main() -> None:
    default = _ROOT / "samples" / "blackthorn-ellis-test-letterhead.docx"
    path = _add_header_footer(default_out=default)
    print(path)


if __name__ == "__main__":
    main()
