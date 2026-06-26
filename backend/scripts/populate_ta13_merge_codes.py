#!/usr/bin/env python3
"""Insert Canary merge codes into the Law Society TA13 completion form."""

from __future__ import annotations

import argparse
import shutil
import sys
from pathlib import Path

_ROOT = Path(__file__).resolve().parents[1]
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))

from docx import Document


def _set_cell_text(cell, text: str) -> None:
    if cell.paragraphs:
        cell.paragraphs[0].text = text
        for extra in cell.paragraphs[1:]:
            extra.clear()
    else:
        cell.text = text


def populate_ta13(src: Path, dest: Path) -> None:
    doc = Document(str(src))

    if len(doc.tables) < 2:
        raise ValueError("Expected TA13 to contain at least two tables (property + lender).")

    prop_table = doc.tables[0]
    lender_table = doc.tables[1]

    _set_cell_text(prop_table.rows[0].cells[1], "[PROPERTY_ADDRESS_BLOCK]")
    _set_cell_text(prop_table.rows[1].cells[1], "[PRIMARY_CLIENT_NAME]")
    _set_cell_text(lender_table.rows[1].cells[0], "[EXISTING_LENDER_NAME]")
    _set_cell_text(lender_table.rows[1].cells[1], "[PROPERTY_CHARGE_DATE]")

    para_replacements: dict[int, str] = {
        20: "4.2 Name of bank:",
        21: "Address of bank: [FIRM_ADDRESS_BLOCK]",
        22: "Branch sort code: [FIRM_CLIENT_BANK_SORT_CODE]",
        23: "Client account name: [FIRM_CLIENT_BANK_ACCOUNT_NAME]",
        24: "Client account number: [FIRM_CLIENT_BANK_ACCOUNT_NUMBER]",
        25: "Please use reference: [CASE_REF]",
        33: "Signed:",
        34: "[FEE_EARNER_SIGNATURE]",
        35: "[FIRM_TRADING_NAME]",
        36: "Dated: [DATE]",
    }

    for idx, text in para_replacements.items():
        if idx >= len(doc.paragraphs):
            raise ValueError(f"Paragraph index {idx} out of range.")
        doc.paragraphs[idx].text = text

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest))


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--input",
        type=Path,
        default=Path.home() / "Downloads" / "TA13.docx",
        help="Source TA13 .docx (default: ~/Downloads/TA13.docx)",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=None,
        help="Output path (default: overwrite --input after backing up to _originals/)",
    )
    args = parser.parse_args()
    src = args.input.expanduser().resolve()
    if not src.is_file():
        raise SystemExit(f"Not found: {src}")

    if args.output:
        dest = args.output.expanduser().resolve()
        populate_from = src
    else:
        backup_dir = src.parent / "_originals"
        backup_dir.mkdir(exist_ok=True)
        backup = backup_dir / src.name
        if not backup.is_file():
            shutil.copy2(src, backup)
        dest = src
        populate_from = backup

    populate_ta13(populate_from, dest)
    print(f"Wrote {dest}")


if __name__ == "__main__":
    main()
