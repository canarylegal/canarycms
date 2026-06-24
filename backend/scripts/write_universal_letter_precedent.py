"""Emit the canonical Letter precedent .docx scaffold used as the reserved BLANK_LETTER template.

Produces the full universal letter shell: recipient names + org/address block, date, refs,
salutation, ``Re:`` subject line, the ``[PRECEDENT_BODY]`` insertion-point marker (where the
chosen precedent's body is spliced in at compose time — see
:func:`app.docx_util.splice_precedent_into_blank_letter`), and the signature block
(``[CONTACT_LETTER_SIGN_OFF]`` / ``[FEE_EARNER]`` / ``[FIRM_TRADING_NAME]``).

Run manually; not invoked by the running app. The canonical seed copy lives at
``backend/precedents_seed/bundle/g0_blank_letter.docx`` — exact vertical spacing in the
generated output may differ from that file (the bundled copy is hand-tuned). Use this script
to regenerate from scratch or as a starting point for custom firm scaffolds.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


def main() -> None:
    out = Path(sys.argv[1]).expanduser().resolve() if len(sys.argv) > 1 else Path("Universal-letter-precedent.docx")

    doc = Document()
    names_row = (
        "[TITLE] [FIRST_INITIAL] [MIDDLE_INITIAL] [LAST_NAME] "
        "[TITLE_2] [FIRST_INITIAL_2] [MIDDLE_INITIAL_2] [LAST_NAME_2] "
        "[TITLE_3] [FIRST_INITIAL_3] [MIDDLE_INITIAL_3] [LAST_NAME_3] "
        "[TITLE_4] [FIRST_INITIAL_4] [MIDDLE_INITIAL_4] [LAST_NAME_4]"
    )
    # Single paragraph so Word does not insert spacing between the names row and org/address block.
    doc.add_paragraph(f"{names_row}\n[ORG_AND_ADDRESS_BLOCK]")
    body = [
        "",
        "[DATE]",
        "",
        "Your Ref: [CONTACT_REF]",
        "Our Ref: [FEE_EARNER_INITIALS]/[CASE_REF]",
        "",
        "[CONTACT_LETTER_DEAR]",
        "",
        "Re: [MATTER_DESCRIPTION]",
        "[SOLICITOR_OUR_CLIENT_LINE]",
        "[SOLICITOR_YOUR_CLIENT_LINE]",
        "",
        "[PRECEDENT_BODY]",
        "",
        "[CONTACT_LETTER_SIGN_OFF]",
        "",
        "",
        "[FEE_EARNER]",
        "[FIRM_TRADING_NAME]",
    ]
    for line in body:
        doc.add_paragraph(line)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
