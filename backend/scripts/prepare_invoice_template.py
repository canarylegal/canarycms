#!/usr/bin/env python3
"""Prepare an invoice template .docx with indexed merge codes (legacy helper).

Prefer :mod:`write_universal_invoice_precedent` for the bundled universal template
(``precedents_seed/bundle/g1_invoice_template.docx``). This script keeps headers/footers
from a source file when you need a firm-branded starting point.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _clear_body_keep_section(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def _right_align_cell(cell) -> None:
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def prepare_invoice_template(src: Path, dest: Path, *, slots: int = 25) -> None:
    doc = Document(str(src))
    _clear_body_keep_section(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("INVOICE")
    run.bold = True
    run.font.size = Pt(16)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run("[INVOICE_NUMBER]  ·  [INVOICE_DATE]")
    meta_run.font.size = Pt(11)

    doc.add_paragraph("")

    bill = doc.add_paragraph()
    bill.add_run("Bill to: ").bold = True
    bill.add_run("[INVOICE_BILL_TO]")

    doc.add_paragraph("Matter: [CASE_REF] — [MATTER_DESCRIPTION]")
    doc.add_paragraph("Fee earner: [FEE_EARNER]")
    doc.add_paragraph("")

    table = doc.add_table(rows=1 + slots, cols=5)
    table.style = "Table Grid"
    headers = ("Type", "Description", "Net", "VAT", "Total")
    hdr = table.rows[0].cells
    for i, label in enumerate(headers):
        hdr[i].text = label
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
            if i >= 2:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for i in range(1, slots + 1):
        tag = f"{i:02d}"
        row = table.rows[i].cells
        row[0].text = f"[INVOICE_{tag}_TYPE]"
        row[1].text = f"[INVOICE_{tag}_DESCRIPTION]"
        row[2].text = f"[INVOICE_{tag}_NET]"
        row[3].text = f"[INVOICE_{tag}_VAT]"
        row[4].text = f"[INVOICE_{tag}_TOTAL]"
        for j in range(2, 5):
            _right_align_cell(row[j])

    doc.add_paragraph("")

    totals = doc.add_table(rows=3, cols=2)
    totals.style = "Table Grid"
    total_rows = (
        ("Net total", "[INVOICE_NET_TOTAL]"),
        ("VAT total", "[INVOICE_VAT_TOTAL]"),
        ("Invoice total", "[INVOICE_TOTAL]"),
    )
    for i, (label, code) in enumerate(total_rows):
        totals.rows[i].cells[0].text = label
        amt_cell = totals.rows[i].cells[1]
        amt_cell.text = code
        _right_align_cell(amt_cell)
        if i == 2:
            for cell in totals.rows[i].cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True

    doc.add_paragraph("")
    footer = doc.add_paragraph("[FIRM_TRADING_NAME]")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(10)

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest))


def main() -> None:
    parser = argparse.ArgumentParser(description="Prepare invoice template with merge codes")
    parser.add_argument("--src", type=Path, required=True, help="Source .docx (letterhead headers/footers kept)")
    parser.add_argument("--dest", type=Path, required=True, help="Output .docx to upload as invoice template")
    parser.add_argument("--slots", type=int, default=25, help="Number of invoice line rows (default 25)")
    args = parser.parse_args()
    prepare_invoice_template(args.src, args.dest, slots=args.slots)
    print(f"Wrote {args.dest}")


if __name__ == "__main__":
    main()
