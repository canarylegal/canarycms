"""Digital letterhead image / media package merging."""

import base64
import io
import zipfile

from docx import Document
from docx.shared import Inches

from app.docx_util import apply_digital_letterhead_headers_footers

_PNG_1PX = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
)


def _letterhead_with_header_image() -> bytes:
    doc = Document()
    doc.add_paragraph("")
    doc.sections[0].header.paragraphs[0].add_run().add_picture(
        io.BytesIO(_PNG_1PX), width=Inches(1)
    )
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _blank_precedent() -> bytes:
    doc = Document()
    doc.add_paragraph("Letter body")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_letterhead_merge_copies_media_and_header_rels() -> None:
    merged = apply_digital_letterhead_headers_footers(_blank_precedent(), _letterhead_with_header_image())
    with zipfile.ZipFile(io.BytesIO(merged)) as z:
        assert "word/media/image1.png" in z.namelist()
        rels = z.read("word/_rels/header1.xml.rels").decode()
        assert "media/image1.png" in rels
        assert 'r:embed="rId1"' in z.read("word/header1.xml").decode()


def test_letterhead_merge_content_types_includes_png() -> None:
    merged = apply_digital_letterhead_headers_footers(_blank_precedent(), _letterhead_with_header_image())
    with zipfile.ZipFile(io.BytesIO(merged)) as z:
        ct = z.read("[Content_Types].xml").decode()
        assert 'Extension="png"' in ct


def _letterhead_with_body_image() -> bytes:
    doc = Document()
    doc.add_paragraph().add_run().add_picture(io.BytesIO(_PNG_1PX), width=Inches(1))
    doc.add_paragraph("")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_letterhead_merge_prepends_body_masthead_image() -> None:
    merged = apply_digital_letterhead_headers_footers(_blank_precedent(), _letterhead_with_body_image())
    with zipfile.ZipFile(io.BytesIO(merged)) as z:
        assert "word/media/image1.png" in z.namelist()
        doc_xml = z.read("word/document.xml").decode()
        assert "Letter body" in doc_xml
        assert "drawing" in doc_xml or "blip" in doc_xml
        rels = z.read("word/_rels/document.xml.rels").decode()
        assert "media/image1.png" in rels


def _letterhead_with_first_page_footer() -> bytes:
    """Mimics Ashbourne and Finch-style letterheads: firm lines on the first-page footer only."""
    doc = Document()
    doc.add_paragraph("")
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True
    sec.first_page_footer.paragraphs[0].add_run("49 Bell Street, Sawbridgeworth")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_letterhead_merge_copies_first_page_footer() -> None:
    merged = apply_digital_letterhead_headers_footers(_blank_precedent(), _letterhead_with_first_page_footer())
    out = Document(io.BytesIO(merged))
    sec = out.sections[0]
    assert sec.different_first_page_header_footer is True
    footer_text = "\n".join(p.text for p in sec.first_page_footer.paragraphs)
    assert "Bell Street" in footer_text


def _letterhead_with_compat_mode(*, mode: str) -> bytes:
    import xml.etree.ElementTree as ET

    doc = Document()
    doc.add_paragraph("")
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True
    sec.first_page_footer.paragraphs[0].add_run("Footer line one")
    buf = io.BytesIO()
    doc.save(buf)
    parts: dict[str, bytes] = {}
    with zipfile.ZipFile(io.BytesIO(buf.getvalue()), "r") as zin:
        for info in zin.infolist():
            parts[info.filename] = zin.read(info.filename)
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    root = ET.fromstring(parts["word/settings.xml"])
    compat = root.find(f"{{{W}}}compat")
    if compat is None:
        compat = ET.SubElement(root, f"{{{W}}}compat")
    for child in list(compat):
        compat.remove(child)
    cs = ET.SubElement(compat, f"{{{W}}}compatSetting")
    cs.set(f"{{{W}}}name", "compatibilityMode")
    cs.set(f"{{{W}}}val", mode)
    parts["word/settings.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in parts.items():
            zout.writestr(name, data)
    return out.getvalue()


def _compat_mode_from_docx(data: bytes) -> str | None:
    import xml.etree.ElementTree as ET

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        root = ET.fromstring(z.read("word/settings.xml"))
    compat = root.find(f"{{{W}}}compat")
    if compat is None:
        return None
    for cs in compat.findall(f"{{{W}}}compatSetting"):
        if cs.get(f"{{{W}}}name") == "compatibilityMode":
            return cs.get(f"{{{W}}}val")
    return None


def test_letterhead_merge_applies_letterhead_compatibility_mode() -> None:
    """Precedent scaffolds often use a newer Word compat mode than the uploaded letterhead."""
    lh = _letterhead_with_compat_mode(mode="11")
    prec = _letterhead_with_compat_mode(mode="14")
    merged = apply_digital_letterhead_headers_footers(prec, lh)
    assert _compat_mode_from_docx(merged) == "11"
