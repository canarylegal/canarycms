"""Quote table row cleanup after merge."""

import io
import re
from pathlib import Path

import pytest
from docx import Document

from app.docx_util import merge_precedent_codes, strip_empty_quote_table_rows


def test_strip_empty_quote_table_rows_removes_unused_slots() -> None:
    src_path = Path("/home/colin/Downloads/Letterhead.docx")
    if not src_path.is_file():
        pytest.skip("Letterhead.docx not available")
    src = src_path.read_bytes()
    fields = {}
    for i in range(1, 16):
        tag = f"{i:02d}"
        fields[f"[QUOTE_{tag}_LABEL]"] = f"Line {i}" if i <= 3 else ""
        fields[f"[QUOTE_{tag}_AMOUNT]"] = f"£{i * 100}.00" if i <= 3 else ""
    fields["[QUOTE_PROPERTY_VALUE]"] = "£150,000.00"
    merged = merge_precedent_codes(src, fields)
    cleaned = strip_empty_quote_table_rows(merged)
    doc = Document(io.BytesIO(cleaned))
    rows = doc.tables[0].rows
    assert len(rows) == 4  # header + 3 data rows
    xml = rows[-1].cells[0].text
    assert "QUOTE_" not in xml
    assert re.search(r"\[QUOTE_", Document(io.BytesIO(cleaned)).tables[0].rows[-1].cells[1].text) is None


def test_quote_merge_fields_clear_unused_slots() -> None:
    from app.compose_quote import _quote_line_merge_fields
    from app.fee_scale_calc import ComputedQuoteLine
    from app.models import FeeScaleLineKind

    lines = [
        ComputedQuoteLine(
            line_id=None,
            name="Legal Fees",
            line_kind=FeeScaleLineKind.item,
            amount_pence=112500,
            editable=True,
            is_bold=False,
            align_right=True,
        ),
    ]
    fields = _quote_line_merge_fields(lines, property_value_pence=150_000_00)
    assert fields["[QUOTE_01_LABEL]"] == "Legal Fees"
    assert fields["[QUOTE_25_LABEL]"] == ""
    assert fields["[QUOTE_25_AMOUNT]"] == ""
