"""Solicitor letter Our/Your client merge lines."""

import io
from types import SimpleNamespace
from uuid import uuid4

from docx import Document

from app.docx_util import (
    PRECEDENT_BODY_MARKER,
    build_merge_fields,
    merge_precedent_codes,
    splice_precedent_into_blank_letter,
)
from app.matter_contact_constants import CLIENT_SLUG, LAWYERS_SLUG


def _contact(*, name: str, matter_type: str, lawyer_client_ids: list | None = None) -> SimpleNamespace:
    return SimpleNamespace(
        id=uuid4(),
        name=name,
        matter_contact_type=matter_type,
        lawyer_client_ids=lawyer_client_ids or [],
        type="organisation",
        trading_name=None,
        company_name=name,
        address_line1="",
        address_line2="",
        city="",
        county="",
        postcode="",
        matter_contact_reference="",
    )


def test_solicitor_lines_filled_when_compose_contact_is_lawyer() -> None:
    our_client = _contact(name="Kieba Homes Ltd", matter_type=CLIENT_SLUG)
    their_client = _contact(name="Buyer Co", matter_type=CLIENT_SLUG)
    lawyer = _contact(
        name="Other Side LLP",
        matter_type=LAWYERS_SLUG,
        lawyer_client_ids=[str(their_client.id)],
    )
    fields = build_merge_fields(
        SimpleNamespace(title="Sale", case_number="CMW/1"),
        lawyer_slots=[(lawyer, [their_client])],
        compose_selected_contact=lawyer,
        ordered_client_contacts=[our_client],
    )
    assert fields["[SOLICITOR_OUR_CLIENT_LINE]"] == "Our Client: Kieba Homes Ltd"
    assert fields["[SOLICITOR_YOUR_CLIENT_LINE]"] == "Your Client: Buyer Co"


def test_solicitor_lines_empty_for_client_letter() -> None:
    client = _contact(name="Mr A Buyer", matter_type=CLIENT_SLUG)
    fields = build_merge_fields(
        SimpleNamespace(title="Sale", case_number="CMW/1"),
        compose_selected_contact=client,
        ordered_client_contacts=[client],
    )
    assert fields["[SOLICITOR_OUR_CLIENT_LINE]"] == ""
    assert fields["[SOLICITOR_YOUR_CLIENT_LINE]"] == ""


def test_splice_strips_duplicate_solicitor_scaffold_from_precedent() -> None:
    blank = Document()
    blank.add_paragraph("Shell")
    blank.add_paragraph(PRECEDENT_BODY_MARKER)
    blank_buf = io.BytesIO()
    blank.save(blank_buf)

    prec = Document()
    for line in [
        "Re: [MATTER_DESCRIPTION]",
        "[SOLICITOR_OUR_CLIENT_LINE]",
        "[SOLICITOR_YOUR_CLIENT_LINE]",
        "",
        "Letter body paragraph",
    ]:
        prec.add_paragraph(line)
    prec_buf = io.BytesIO()
    prec.save(prec_buf)

    merged = splice_precedent_into_blank_letter(blank_buf.getvalue(), prec_buf.getvalue())
    out = Document(io.BytesIO(merged))
    texts = [p.text for p in out.paragraphs]
    assert texts.count("Re: [MATTER_DESCRIPTION]") == 0
    assert "Letter body paragraph" in texts


def test_solicitor_line_paragraphs_removed_when_empty() -> None:
    doc = Document()
    doc.add_paragraph("[SOLICITOR_OUR_CLIENT_LINE]")
    doc.add_paragraph("[SOLICITOR_YOUR_CLIENT_LINE]")
    doc.add_paragraph("Body")
    buf = io.BytesIO()
    doc.save(buf)
    merged = merge_precedent_codes(
        buf.getvalue(),
        {
            "[SOLICITOR_OUR_CLIENT_LINE]": "",
            "[SOLICITOR_YOUR_CLIENT_LINE]": "",
        },
    )
    out = Document(io.BytesIO(merged))
    assert [p.text.strip() for p in out.paragraphs if p.text.strip()] == ["Body"]
