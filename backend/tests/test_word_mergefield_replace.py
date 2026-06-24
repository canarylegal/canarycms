"""Word MERGEFIELD → Canary [CODE] conversion."""

import io
import zipfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.docx_util import replace_word_mergefields_in_docx_bytes


def _docx_with_fldsimple_mergefield(field_name: str) -> bytes:
    doc = Document()
    para = doc.add_paragraph()
    para.clear()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), f" MERGEFIELD {field_name} ")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "cached"
    r.append(t)
    fld.append(r)
    para._p.append(fld)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_with_complex_mergefield(field_name: str) -> bytes:
    doc = Document()
    para = doc.add_paragraph()
    para.clear()
    p = para._p

    def add_run(children: list) -> None:
        r = OxmlElement("w:r")
        for child in children:
            r.append(child)
        p.append(r)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    add_run([begin])

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" MERGEFIELD {field_name} "
    add_run([instr])

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    add_run([separate])

    t = OxmlElement("w:t")
    t.text = "cached"
    add_run([t])

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    add_run([end])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_replace_fldsimple_mergefield() -> None:
    merged = replace_word_mergefields_in_docx_bytes(
        _docx_with_fldsimple_mergefield("MATTER__Person_Acting_Full_Name"),
        {"MATTER__Person_Acting_Full_Name": "[FEE_EARNER]"},
    )
    out = Document(io.BytesIO(merged))
    assert out.paragraphs[0].text == "[FEE_EARNER]"
    assert "MERGEFIELD" not in zipfile.ZipFile(io.BytesIO(merged)).read("word/document.xml").decode()


def test_replace_complex_mergefield() -> None:
    merged = replace_word_mergefields_in_docx_bytes(
        _docx_with_complex_mergefield("CONVEYANCING_D__Price_Total"),
        {"CONVEYANCING_D__Price_Total": "[QUOTE_PROPERTY_VALUE]"},
    )
    out = Document(io.BytesIO(merged))
    assert "[QUOTE_PROPERTY_VALUE]" in out.paragraphs[0].text
    assert "MERGEFIELD" not in zipfile.ZipFile(io.BytesIO(merged)).read("word/document.xml").decode()
