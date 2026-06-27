"""Quote compose uses built-in template (not blank letter) when no quote letterhead."""

import io
import zipfile

from app.docx_util import write_quote_template_docx_bytes
def test_builtin_quote_template_has_fee_table_slots() -> None:
    raw = write_quote_template_docx_bytes(slots=3)
    xml = zipfile.ZipFile(io.BytesIO(raw)).read("word/document.xml").decode()
    assert "[QUOTE_01_LABEL]" in xml
    assert "[QUOTE_03_VAT]" in xml
    assert "[PRECEDENT_BODY]" not in xml


def test_load_quote_template_without_letterhead_uses_builtin() -> None:
    raw = write_quote_template_docx_bytes()
    xml = zipfile.ZipFile(io.BytesIO(raw)).read("word/document.xml").decode()
    assert "[QUOTE_01_LABEL]" in xml
    assert "[QUOTE_01_AMOUNT]" in xml


def test_builtin_quote_template_merges_fee_lines() -> None:
    from app.docx_util import apply_quote_table_presentation, merge_precedent_codes, strip_empty_quote_table_rows
    from app.compose_quote import _quote_line_merge_fields
    from app.models import FeeScaleLineKind
    from app.fee_scale_calc import ComputedQuoteLine

    raw = write_quote_template_docx_bytes(slots=5)
    lines = [
        ComputedQuoteLine(
            line_id=None,
            name="Legal fee",
            line_kind=FeeScaleLineKind.item,
            amount_pence=92500,
            vat_pence=18500,
            editable=False,
            is_bold=False,
            align_right=True,
        ),
    ]
    fields = _quote_line_merge_fields(lines, property_value_pence=15000000)
    merged = merge_precedent_codes(raw, fields)
    cleaned = strip_empty_quote_table_rows(merged)
    styled = apply_quote_table_presentation(cleaned, lines)
    xml = zipfile.ZipFile(io.BytesIO(styled)).read("word/document.xml").decode()
    assert "Legal fee" in xml
    assert "925.00" in xml
    assert "QUOTE_01_LABEL" not in xml
    assert 'w:fill="D0DAEA"' in xml


def test_apply_quote_table_presentation_styles_section_and_total_rows() -> None:
    from app.docx_util import apply_quote_table_presentation, merge_precedent_codes, strip_empty_quote_table_rows
    from app.compose_quote import _quote_line_merge_fields
    from app.models import FeeScaleLineKind
    from app.fee_scale_calc import ComputedQuoteLine

    raw = write_quote_template_docx_bytes(slots=8)
    lines = [
        ComputedQuoteLine(
            line_id=None,
            name="Disbursements/Fees paid to Third Parties",
            line_kind=FeeScaleLineKind.section_header,
            amount_pence=None,
            editable=False,
            is_bold=True,
            align_right=False,
        ),
        ComputedQuoteLine(
            line_id=None,
            name="Telegraphic Transfer Fee",
            line_kind=FeeScaleLineKind.item,
            amount_pence=4000,
            vat_pence=None,
            editable=False,
            is_bold=False,
            align_right=True,
        ),
        ComputedQuoteLine(
            line_id=None,
            name="TOTAL DISBURSEMENTS",
            line_kind=FeeScaleLineKind.subtotal,
            amount_pence=4000,
            vat_pence=None,
            editable=False,
            is_bold=True,
            align_right=True,
        ),
    ]
    fields = _quote_line_merge_fields(lines, property_value_pence=None)
    merged = merge_precedent_codes(raw, fields)
    cleaned = strip_empty_quote_table_rows(merged)
    styled = apply_quote_table_presentation(cleaned, lines)
    xml = zipfile.ZipFile(io.BytesIO(styled)).read("word/document.xml").decode()
    assert "Disbursements/Fees paid to Third Parties" in xml
    assert 'w:fill="EEF2F8"' in xml
    assert "TOTAL DISBURSEMENTS" in xml


def test_apply_quote_table_presentation_preserves_table_widths() -> None:
    """Compose styling must not override template table widths."""
    import io

    from docx import Document
    from docx.oxml.ns import qn

    from app.compose_quote import _quote_line_merge_fields
    from app.docx_util import (
        _docx_set_table_width_pct,
        apply_quote_table_presentation,
        merge_precedent_codes,
        strip_empty_quote_table_rows,
    )
    from app.fee_scale_calc import ComputedQuoteLine
    from app.models import FeeScaleLineKind

    doc = Document()
    fee = doc.add_table(rows=3, cols=3)
    _docx_set_table_width_pct(fee, 4500)
    fee.rows[0].cells[0].text = "Description"
    fee.rows[0].cells[1].text = "Amount"
    fee.rows[0].cells[2].text = "VAT"
    for i in range(1, 3):
        tag = f"{i:02d}"
        fee.rows[i].cells[0].text = f"[QUOTE_{tag}_LABEL]"
        fee.rows[i].cells[1].text = f"[QUOTE_{tag}_AMOUNT]"
        fee.rows[i].cells[2].text = f"[QUOTE_{tag}_VAT]"

    summary = doc.add_table(rows=1, cols=2)
    _docx_set_table_width_pct(summary, 4500)
    summary.rows[0].cells[0].text = "Estimate grand total:"
    summary.rows[0].cells[1].text = "[QUOTE_GRAND_TOTAL]"

    buf = io.BytesIO()
    doc.save(buf)
    raw = buf.getvalue()

    def _table_layout(doc_bytes: bytes, index: int) -> dict[str, object]:
        loaded = Document(io.BytesIO(doc_bytes))
        tbl = loaded.tables[index]._tbl
        tbl_pr = tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW")) if tbl_pr is not None else None
        layout = tbl_pr.find(qn("w:tblLayout")) if tbl_pr is not None else None
        grid = tbl.find(qn("w:tblGrid"))
        cols = [gc.get(qn("w:w")) for gc in grid.findall(qn("w:gridCol"))] if grid is not None else []
        duplicate_tbl_w = len(tbl_pr.findall(qn("w:tblW"))) if tbl_pr is not None else 0
        return {
            "tblW": (tbl_w.get(qn("w:w")), tbl_w.get(qn("w:type"))) if tbl_w is not None else None,
            "layout": layout.get(qn("w:type")) if layout is not None else None,
            "gridCol": cols,
            "duplicate_tblW": duplicate_tbl_w,
        }

    before_fee = _table_layout(raw, 0)
    before_summary = _table_layout(raw, 1)

    lines = [
        ComputedQuoteLine(
            line_id=None,
            name="Legal fee",
            line_kind=FeeScaleLineKind.item,
            amount_pence=92500,
            vat_pence=18500,
            editable=False,
            is_bold=False,
            align_right=True,
        ),
    ]
    fields = _quote_line_merge_fields(lines, property_value_pence=15000000)
    merged = merge_precedent_codes(raw, fields)
    cleaned = strip_empty_quote_table_rows(merged)
    styled = apply_quote_table_presentation(cleaned, lines)

    after_fee = _table_layout(styled, 0)
    after_summary = _table_layout(styled, 1)

    assert before_fee["tblW"] == ("4500", "pct")
    assert before_summary["tblW"] == ("4500", "pct")
    assert after_fee["tblW"] == before_fee["tblW"]
    assert after_summary["tblW"] == before_fee["tblW"]
    assert after_fee["duplicate_tblW"] == 1
    assert after_summary["duplicate_tblW"] == 1
    assert after_fee["layout"] == "fixed"
    assert after_summary["layout"] == "fixed"
    assert after_summary["gridCol"] == [
        after_fee["gridCol"][0],
        str(int(after_fee["gridCol"][1]) + int(after_fee["gridCol"][2])),
    ]
