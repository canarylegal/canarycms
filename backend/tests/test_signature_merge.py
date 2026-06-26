"""Fee earner signature image merge code."""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

from docx import Document

from app.docx_util import _empty_precedent_field_map, inject_merge_code_images, merge_precedent_codes

# Minimal valid 1x1 PNG.
_TINY_PNG = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01"
    b"\x00\x00\x05\x00\x01\r\n-\xdb\x00\x00\x00\x00IEND\xaeB`\x82"
)


def _docx_with_signature_code(*, trailing_nbsp: bool = False) -> bytes:
    doc = Document()
    text = "[FEE_EARNER_SIGNATURE]"
    if trailing_nbsp:
        text += "\xa0"
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_merge_precedent_codes_preserves_signature_placeholder() -> None:
    fields = _empty_precedent_field_map()
    fields["[MATTER_DESCRIPTION]"] = "Sale of 1 High Street"
    merged = merge_precedent_codes(_docx_with_signature_code(), fields)
    doc = Document(io.BytesIO(merged))
    assert doc.paragraphs[0].text.strip() == "[FEE_EARNER_SIGNATURE]"


def test_inject_merge_code_images_inserts_picture(tmp_path: Path) -> None:
    img = tmp_path / "sig.png"
    img.write_bytes(_TINY_PNG)
    merged = merge_precedent_codes(_docx_with_signature_code(trailing_nbsp=True), _empty_precedent_field_map())
    out = inject_merge_code_images(merged, {"[FEE_EARNER_SIGNATURE]": img})

    doc = Document(io.BytesIO(out))
    assert doc.paragraphs[0].text.strip() == ""
    assert doc.paragraphs[0].runs[0]._element.xpath(".//a:blip")  # inline image present

    with zipfile.ZipFile(io.BytesIO(out)) as zf:
        media = [n for n in zf.namelist() if n.startswith("word/media/")]
        assert media
