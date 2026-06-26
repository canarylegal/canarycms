"""Merge-code formatting modifiers ([b:CODE], [i:CODE], etc.)."""

import io

import pytest
from docx import Document

from app.docx_util import (
    _MergeTextSegment,
    _parse_modifier_letters,
    _replace_merge_tokens_in_ooxml_text,
    _replace_merge_tokens_to_segments,
    _segments_plain_text,
    merge_precedent_codes,
)


def test_parse_modifier_letters() -> None:
    assert _parse_modifier_letters(None) == (None, None, None)
    assert _parse_modifier_letters("b") == (True, None, None)
    assert _parse_modifier_letters("i") == (None, True, None)
    assert _parse_modifier_letters("u") == (None, None, True)
    assert _parse_modifier_letters("bi") == (True, True, None)
    assert _parse_modifier_letters("biu") == (True, True, True)


def test_replace_merge_tokens_to_segments_plain_and_modified() -> None:
    fields = {"[MATTER_DESCRIPTION]": "Purchase", "[LAST_NAME]": "Smith"}
    segments = _replace_merge_tokens_to_segments(
        "Matter [b:MATTER_DESCRIPTION] for [LAST_NAME] and [i:LAST_NAME]",
        fields,
    )
    assert _segments_plain_text(segments) == "Matter Purchase for Smith and Smith"
    assert segments[0].text == "Matter "
    assert segments[0].bold is None
    assert segments[1].text == "Purchase"
    assert segments[1].bold is True
    assert segments[2].text == " for "
    assert segments[3].text == "Smith"
    assert segments[3].bold is None
    assert segments[4].text == " and "
    assert segments[5].text == "Smith"
    assert segments[5].italic is True


def test_replace_merge_tokens_unknown_code_left_literal() -> None:
    fields = {"[MATTER_DESCRIPTION]": "X"}
    segments = _replace_merge_tokens_to_segments("[b:NOT_IN_FIELDS]", fields)
    assert segments == [_MergeTextSegment("[b:NOT_IN_FIELDS]")]


def test_ooxml_text_replace_strips_modifiers() -> None:
    fields = {"[MATTER_DESCRIPTION]": "Bold &amp; Co"}
    out = _replace_merge_tokens_in_ooxml_text("See [b:MATTER_DESCRIPTION] today", fields)
    assert out == "See Bold &amp; Co today"


def _docx_bytes_with_paragraph(text: str) -> bytes:
    doc = Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_merge_precedent_codes_applies_bold_in_docx() -> None:
    src = _docx_bytes_with_paragraph("Reference: [b:MATTER_DESCRIPTION]")
    merged = merge_precedent_codes(src, {"[MATTER_DESCRIPTION]": "Bold matter"})
    out = Document(io.BytesIO(merged))
    para = out.paragraphs[0]
    assert para.text == "Reference: Bold matter"
    bold_runs = [r for r in para.runs if r.text and r.bold]
    assert any("Bold matter" in r.text for r in bold_runs)


def test_merge_precedent_codes_bold_multiline_org_block_single_run() -> None:
    src = _docx_bytes_with_paragraph("[b:ORG_AND_ADDRESS_BLOCK]")
    merged = merge_precedent_codes(
        src,
        {"[ORG_AND_ADDRESS_BLOCK]": "Acme Ltd\n10 High Street\nLondon"},
    )
    out = Document(io.BytesIO(merged))
    para = out.paragraphs[0]
    assert para.text == "Acme Ltd\n10 High Street\nLondon"
    text_runs = [r for r in para.runs if r.text and r.text != "\n"]
    assert len(text_runs) == 1
    assert text_runs[0].bold is True
    assert "Acme Ltd" in text_runs[0].text
    assert "London" in text_runs[0].text


def test_coalesce_split_modifier_token_across_runs() -> None:
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("[b:ORG_")
    p.add_run("AND_ADDRESS_BLOCK]")
    buf = io.BytesIO()
    doc.save(buf)
    merged = merge_precedent_codes(buf.getvalue(), {"[ORG_AND_ADDRESS_BLOCK]": "Bold block"})
    out = Document(io.BytesIO(merged))
    assert out.paragraphs[0].text == "Bold block"
    assert any(r.bold for r in out.paragraphs[0].runs if r.text)


def test_coalesce_split_token_preserves_bold_prefix_run() -> None:
    """ONLYOFFICE often splits ``[b:CODE]`` across runs; bold ``Re:`` prefix must survive."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    p = doc.add_paragraph()
    for r in list(p._p.findall(qn("w:r"))):
        p._p.remove(r)

    def add_run(text: str, *, bold: bool = False) -> None:
        r = OxmlElement("w:r")
        if bold:
            r_pr = OxmlElement("w:rPr")
            r_pr.append(OxmlElement("w:b"))
            r_pr.append(OxmlElement("w:bCs"))
            r.append(r_pr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        p._p.append(r)

    add_run("Re:", bold=True)
    add_run(" [")
    add_run("b:")
    add_run("MATTER_DESCRIPTION]")
    buf = io.BytesIO()
    doc.save(buf)
    merged = merge_precedent_codes(buf.getvalue(), {"[MATTER_DESCRIPTION]": "Sale of 1 High Street"})
    out = Document(io.BytesIO(merged))
    para = out.paragraphs[0]
    assert para.text == "Re: Sale of 1 High Street"
    bold_runs = [r for r in para.runs if r.text and r.bold]
    assert any(r.text.startswith("Re:") for r in bold_runs)
    assert any("Sale of 1 High Street" in r.text and r.bold for r in para.runs if r.text)


def test_merge_precedent_codes_plain_code_unstyled() -> None:
    src = _docx_bytes_with_paragraph("[MATTER_DESCRIPTION]")
    merged = merge_precedent_codes(src, {"[MATTER_DESCRIPTION]": "Plain"})
    out = Document(io.BytesIO(merged))
    para = out.paragraphs[0]
    assert para.text == "Plain"
    assert not any(r.bold for r in para.runs if r.text)


def test_merge_precedent_codes_preserves_static_bold_in_same_paragraph() -> None:
    doc = Document()
    para = doc.add_paragraph()
    bold = para.add_run("Important notice: ")
    bold.bold = True
    para.add_run("Ref [MATTER_DESCRIPTION]")
    buf = io.BytesIO()
    doc.save(buf)
    merged = merge_precedent_codes(buf.getvalue(), {"[MATTER_DESCRIPTION]": "Purchase"})
    out = Document(io.BytesIO(merged))
    para = out.paragraphs[0]
    assert para.text == "Important notice: Ref Purchase"
    bold_runs = [r for r in para.runs if r.text and r.bold]
    assert any("Important notice:" in r.text for r in bold_runs)


def test_merge_precedent_codes_preserves_bold_only_paragraph() -> None:
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("Terms and conditions apply.")
    run.bold = True
    buf = io.BytesIO()
    doc.save(buf)
    merged = merge_precedent_codes(buf.getvalue(), {"[MATTER_DESCRIPTION]": "Unused"})
    out = Document(io.BytesIO(merged))
    para = out.paragraphs[0]
    assert para.text == "Terms and conditions apply."
    assert any(r.bold for r in para.runs if r.text)
