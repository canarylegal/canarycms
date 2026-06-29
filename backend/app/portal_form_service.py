"""Portal form templates (precedents) and client submissions."""

from __future__ import annotations

import re
import uuid
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from fastapi import HTTPException, UploadFile, status
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.alert_dispatch import AlertKind, dispatch_alert, firm_alerts_configured, portal_public_url
from app.audit import log_event
from app.fee_scale_service import fee_scale_matches_case
from app.file_storage import FILES_ROOT, case_file_paths, ensure_files_root
from app.models import (
    Case,
    CaseContact,
    Contact,
    ContactPortalAccess,
    File,
    FileCategory,
    PortalFormFieldType,
    PortalFormSubmission,
    PortalFormSubmissionStatus,
    PortalFormTemplate,
    PortalFormTemplateField,
    User,
)
from app.portal_activity import log_portal_activity
from app.portal_notifications import notify_portal_staff_form_completed
from app.portal_service import client_matter_description, contact_display_name, portal_access_is_active, resolve_matter_contact_email
from app.portal_case import require_case_portal_enabled
from app.portal_notifications import ALERTS_NOT_CONFIGURED_MSG
from app.quote_portal_service import pick_portal_grant_for_case

_FIELD_KEY_RE = re.compile(r"^[a-z][a-z0-9_]{0,79}$")

FORM_EMAIL_SEND_FAILED_MSG = (
    "The notification e-mail could not be delivered. The form is still available on the portal."
)


def _form_email_skip_reason(db: Session, *, email_sent: bool) -> str | None:
    if email_sent:
        return None
    if not firm_alerts_configured(db):
        return ALERTS_NOT_CONFIGURED_MSG
    return FORM_EMAIL_SEND_FAILED_MSG


def utcnow() -> datetime:
    return datetime.now(timezone.utc)


def _normalize_field_key(key: str) -> str:
    k = (key or "").strip().lower()
    if not _FIELD_KEY_RE.match(k):
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Field key must start with a letter and use lowercase letters, digits, or underscores.",
        )
    return k


def template_matches_case(
    template: PortalFormTemplate,
    *,
    matter_head_type_id: uuid.UUID | None,
    matter_sub_type_id: uuid.UUID | None,
) -> bool:
    return fee_scale_matches_case(
        template,  # same scope columns as FeeScale
        matter_head_type_id=matter_head_type_id,
        matter_sub_type_id=matter_sub_type_id,
    )


def list_templates_for_case(db: Session, case_id: uuid.UUID) -> list[PortalFormTemplate]:
    case = db.get(Case, case_id)
    if case is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Case not found")
    rows = db.execute(select(PortalFormTemplate).order_by(PortalFormTemplate.name.asc())).scalars().all()
    return [t for t in rows if template_matches_case(t, matter_head_type_id=case.matter_head_type_id, matter_sub_type_id=case.matter_sub_type_id)]


def load_template_fields(db: Session, template_id: uuid.UUID) -> list[PortalFormTemplateField]:
    return (
        db.execute(
            select(PortalFormTemplateField)
            .where(PortalFormTemplateField.template_id == template_id)
            .order_by(PortalFormTemplateField.sort_order, PortalFormTemplateField.label)
        )
        .scalars()
        .all()
    )


def _active_pending(
    db: Session,
    *,
    case_id: uuid.UUID,
    template_id: uuid.UUID,
    contact_id: uuid.UUID,
) -> PortalFormSubmission | None:
    return db.execute(
        select(PortalFormSubmission)
        .where(
            PortalFormSubmission.case_id == case_id,
            PortalFormSubmission.template_id == template_id,
            PortalFormSubmission.contact_id == contact_id,
            PortalFormSubmission.status == PortalFormSubmissionStatus.pending,
        )
        .limit(1)
    ).scalar_one_or_none()


def form_submission_file_list_item(db: Session, submission: PortalFormSubmission) -> dict[str, Any]:
    contact = db.get(Contact, submission.contact_id)
    template = db.get(PortalFormTemplate, submission.template_id)
    return {
        "id": str(submission.id),
        "status": submission.status.value,
        "contact_name": contact_display_name(contact) if contact else "Contact",
        "template_name": template.name if template else "",
        "sent_at": submission.sent_at.isoformat(),
        "completed_at": submission.completed_at.isoformat() if submission.completed_at else None,
    }


def _form_list_filename(*, template: PortalFormTemplate, contact: Contact) -> str:
    return f"{template.name} — {contact_display_name(contact)}.docx"


def _create_form_list_file(
    db: Session,
    *,
    case: Case,
    template: PortalFormTemplate,
    contact: Contact,
    owner: User,
    body: bytes | None = None,
    awaiting_completion: bool = False,
) -> File:
    if body is None:
        doc = Document()
        doc.add_heading(template.name, level=1)
        if awaiting_completion:
            doc.add_paragraph(
                f"Sent to {contact_display_name(contact)} — awaiting completion via the client portal."
            )
        bio = BytesIO()
        doc.save(bio)
        body = bio.getvalue()

    ensure_files_root()
    file_id = uuid.uuid4()
    filename = _form_list_filename(template=template, contact=contact)
    paths = case_file_paths(case_id=case.id, file_id=file_id, original_filename=filename, folder_path="")
    paths.abs_path.write_bytes(body)
    now = datetime.utcnow()
    row = File(
        id=file_id,
        case_id=case.id,
        owner_id=owner.id,
        category=FileCategory.case_document,
        storage_path=paths.rel_path,
        folder_path=paths.folder_path,
        is_pinned=False,
        original_filename=filename,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        size_bytes=len(body),
        version=1,
        checksum=None,
        parent_file_id=None,
        uploaded_via_portal=False,
        created_at=now,
        updated_at=now,
    )
    db.add(row)
    db.flush()
    return row


def send_form_to_contact(
    db: Session,
    *,
    case_id: uuid.UUID,
    template_id: uuid.UUID,
    contact_id: uuid.UUID,
    actor: User,
) -> tuple[PortalFormSubmission, bool, str | None]:
    require_case_portal_enabled(db, case_id)

    case = db.get(Case, case_id)
    if case is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Case not found")
    template = db.get(PortalFormTemplate, template_id)
    if template is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Form template not found")
    if not template_matches_case(template, matter_head_type_id=case.matter_head_type_id, matter_sub_type_id=case.matter_sub_type_id):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Form template is not available for this matter type")

    on_case = db.execute(
        select(CaseContact).where(CaseContact.case_id == case_id, CaseContact.contact_id == contact_id)
    ).scalar_one_or_none()
    if on_case is None:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Contact is not on this matter")

    contact = db.get(Contact, contact_id)
    if contact is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Contact not found")
    email = resolve_matter_contact_email(db, case_id=case_id, contact_id=contact_id)

    access = db.execute(select(ContactPortalAccess).where(ContactPortalAccess.contact_id == contact_id)).scalar_one_or_none()
    if access is None or not portal_access_is_active(access):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Contact does not have active portal access")

    # Forms are delivery-scoped — folder grants are optional (used for portal navigation when present).
    grant = pick_portal_grant_for_case(db, case_id=case_id, contact_id=contact_id)

    existing = _active_pending(db, case_id=case_id, template_id=template_id, contact_id=contact_id)
    if existing is not None:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="This contact already has a pending form of this type. Void it before sending again.",
        )

    fields = load_template_fields(db, template_id)
    if not fields:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Form template has no fields")

    now = utcnow()
    list_file = _create_form_list_file(
        db,
        case=case,
        template=template,
        contact=contact,
        owner=actor,
        awaiting_completion=True,
    )

    submission = PortalFormSubmission(
        id=uuid.uuid4(),
        case_id=case_id,
        template_id=template_id,
        contact_id=contact_id,
        grant_id=grant.id if grant else None,
        sent_by_user_id=actor.id,
        status=PortalFormSubmissionStatus.pending,
        responses={},
        snapshot_file_id=list_file.id,
        sent_at=now,
    )
    db.add(submission)
    db.flush()

    portal_url = portal_public_url().rstrip("/")
    email_sent = dispatch_alert(
        db,
        AlertKind.portal_form_sent,
        to_email=email,
        context={
            "contact_name": contact_display_name(contact),
            "form_name": template.name,
            "matter_label": client_matter_description(case),
            "portal_url": portal_url,
        },
        actor_user_id=actor.id,
    )
    skip_reason = _form_email_skip_reason(db, email_sent=email_sent)
    log_portal_activity(
        db,
        case_id=case_id,
        contact_id=contact_id,
        grant_id=grant.id if grant else None,
        action="portal.form.sent",
        summary=f"{template.name} sent to {contact_display_name(contact)} via portal",
    )
    log_event(
        db,
        actor_user_id=actor.id,
        action="portal.form.sent",
        entity_type="portal_form_submission",
        entity_id=str(submission.id),
        meta={
            "case_id": str(case_id),
            "template_id": str(template_id),
            "contact_id": str(contact_id),
            "email_sent": email_sent,
        },
    )
    return submission, email_sent, skip_reason


def void_submission(db: Session, *, submission: PortalFormSubmission, actor: User) -> PortalFormSubmission:
    if submission.status in (PortalFormSubmissionStatus.voided, PortalFormSubmissionStatus.superseded):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Form submission is already voided")
    submission.status = PortalFormSubmissionStatus.voided
    submission.voided_at = utcnow()
    db.add(submission)
    log_event(
        db,
        actor_user_id=actor.id,
        action="portal.form.voided",
        entity_type="portal_form_submission",
        entity_id=str(submission.id),
        meta={"case_id": str(submission.case_id)},
    )
    return submission


def list_pending_for_contact(db: Session, contact_id: uuid.UUID) -> list[PortalFormSubmission]:
    return (
        db.execute(
            select(PortalFormSubmission)
            .where(
                PortalFormSubmission.contact_id == contact_id,
                PortalFormSubmission.status == PortalFormSubmissionStatus.pending,
            )
            .order_by(PortalFormSubmission.sent_at.desc())
        )
        .scalars()
        .all()
    )


def get_submission_for_contact(db: Session, submission_id: uuid.UUID, contact_id: uuid.UUID) -> PortalFormSubmission:
    row = db.get(PortalFormSubmission, submission_id)
    if row is None or row.contact_id != contact_id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Form not found")
    return row


def _normalize_select_options(raw: Any) -> list[str]:
    if not raw:
        return []
    if not isinstance(raw, list):
        return []
    out: list[str] = []
    for item in raw:
        s = str(item).strip()
        if not s:
            continue
        if s in out:
            continue
        out.append(s[:200])
        if len(out) >= 50:
            break
    return out


def _field_select_options(field: PortalFormTemplateField) -> list[str]:
    return _normalize_select_options(getattr(field, "select_options", None) or [])


def _validate_responses(fields: list[PortalFormTemplateField], responses: dict[str, Any]) -> dict[str, Any]:
    clean: dict[str, Any] = {}
    for field in fields:
        if field.field_type == PortalFormFieldType.section:
            continue
        key = field.field_key
        raw = responses.get(key)
        if raw is None or (isinstance(raw, str) and not raw.strip()):
            if field.required:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail=f"Required field missing: {field.label}",
                )
            continue
        if field.field_type == PortalFormFieldType.select:
            options = _field_select_options(field)
            if not options:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"Dropdown field has no options: {field.label}",
                )
            if isinstance(raw, bool):
                raw = "Yes" if raw else "No"
            val = str(raw).strip()
            if val not in options:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail=f"Invalid choice for {field.label}",
                )
            clean[key] = val
        elif field.field_type == PortalFormFieldType.file:
            if not isinstance(raw, dict) or not raw.get("file_id"):
                if field.required:
                    raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail=f"File required: {field.label}")
                continue
            clean[key] = {"file_id": str(raw["file_id"]), "filename": str(raw.get("filename") or "upload")}
        else:
            clean[key] = str(raw).strip()
    return clean


def _render_submission_docx(
    *,
    template: PortalFormTemplate,
    fields: list[PortalFormTemplateField],
    responses: dict[str, Any],
    case: Case,
    contact: Contact,
    completed_at: datetime,
) -> bytes:
    doc = Document()
    doc.add_heading(template.name, level=1)
    doc.add_paragraph(f"Matter: {client_matter_description(case)}")
    doc.add_paragraph(f"Contact: {contact_display_name(contact)}")
    doc.add_paragraph(f"Completed: {completed_at.strftime('%d/%m/%Y %H:%M UTC')}")
    doc.add_paragraph("")
    for field in fields:
        if field.field_type == PortalFormFieldType.section:
            doc.add_heading(field.label, level=2)
            continue
        val = responses.get(field.field_key)
        if field.field_type == PortalFormFieldType.file:
            if isinstance(val, dict):
                display = str(val.get("filename") or val.get("file_id") or "—")
            else:
                display = "—"
        else:
            display = str(val) if val not in (None, "") else "—"
        doc.add_paragraph(f"{field.label}: {display}")
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _save_snapshot_on_matter(
    db: Session,
    *,
    submission: PortalFormSubmission,
    template: PortalFormTemplate,
    fields: list[PortalFormTemplateField],
    case: Case,
    contact: Contact,
    owner: User,
) -> File:
    completed_at = submission.completed_at or utcnow()
    data = _render_submission_docx(
        template=template,
        fields=fields,
        responses=submission.responses or {},
        case=case,
        contact=contact,
        completed_at=completed_at,
    )
    existing = db.get(File, submission.snapshot_file_id) if submission.snapshot_file_id else None
    if existing is not None and existing.case_id == case.id:
        paths = case_file_paths(
            case_id=case.id,
            file_id=existing.id,
            original_filename=existing.original_filename,
            folder_path=existing.folder_path or "",
        )
        paths.abs_path.write_bytes(data)
        existing.size_bytes = len(data)
        existing.updated_at = datetime.utcnow()
        db.add(existing)
        db.flush()
        return existing

    row = _create_form_list_file(
        db,
        case=case,
        template=template,
        contact=contact,
        owner=owner,
        body=data,
    )
    return row


async def upload_submission_file(
    db: Session,
    *,
    submission: PortalFormSubmission,
    field_key: str,
    upload: UploadFile,
    contact: Contact,
) -> dict[str, str]:
    if submission.status != PortalFormSubmissionStatus.pending:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Form is not pending")
    template = db.get(PortalFormTemplate, submission.template_id)
    if template is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Template not found")
    fields = load_template_fields(db, template.id)
    field = next((f for f in fields if f.field_key == field_key), None)
    if field is None or field.field_type != PortalFormFieldType.file:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid file field")

    raw = await upload.read()
    if not raw:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Empty file")
    if len(raw) > 25 * 1024 * 1024:
        raise HTTPException(status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE, detail="File too large (max 25 MB)")

    original = Path(upload.filename or "upload").name
    mime = upload.content_type or "application/octet-stream"
    ensure_files_root()
    file_id = uuid.uuid4()
    parent_id = submission.snapshot_file_id
    paths = case_file_paths(
        case_id=submission.case_id,
        file_id=file_id,
        original_filename=original,
        folder_path="",
    )
    paths.abs_path.write_bytes(raw)
    owner = db.get(User, submission.sent_by_user_id) if submission.sent_by_user_id else None
    if owner is None:
        case = db.get(Case, submission.case_id)
        owner = db.get(User, case.fee_earner_user_id) if case else None
    if owner is None:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Could not resolve file owner")

    row = File(
        id=file_id,
        case_id=submission.case_id,
        owner_id=owner.id,
        category=FileCategory.case_document,
        storage_path=paths.rel_path,
        folder_path=paths.folder_path,
        is_pinned=False,
        original_filename=original,
        mime_type=mime,
        size_bytes=len(raw),
        version=1,
        checksum=None,
        parent_file_id=parent_id,
        uploaded_via_portal=True,
        created_at=datetime.utcnow(),
        updated_at=datetime.utcnow(),
    )
    db.add(row)
    db.flush()

    responses = dict(submission.responses or {})
    responses[field_key] = {"file_id": str(file_id), "filename": original}
    submission.responses = responses
    db.add(submission)
    db.flush()
    return {"file_id": str(file_id), "filename": original}


def complete_submission(
    db: Session,
    *,
    submission: PortalFormSubmission,
    contact: Contact,
    responses_in: dict[str, Any],
) -> PortalFormSubmission:
    if submission.status != PortalFormSubmissionStatus.pending:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Form is not pending")

    template = db.get(PortalFormTemplate, submission.template_id)
    if template is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Template not found")
    fields = load_template_fields(db, template.id)
    merged = dict(submission.responses or {})
    merged.update(responses_in)
    validated = _validate_responses(fields, merged)

    case = db.get(Case, submission.case_id)
    if case is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Case not found")
    owner = db.get(User, submission.sent_by_user_id) if submission.sent_by_user_id else db.get(User, case.fee_earner_user_id)
    if owner is None:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Could not resolve file owner")

    now = utcnow()
    submission.responses = validated
    submission.completed_at = now
    submission.status = PortalFormSubmissionStatus.completed
    snapshot = _save_snapshot_on_matter(
        db,
        submission=submission,
        template=template,
        fields=fields,
        case=case,
        contact=contact,
        owner=owner,
    )
    submission.snapshot_file_id = snapshot.id
    db.add(submission)

    log_portal_activity(
        db,
        case_id=submission.case_id,
        contact_id=contact.id,
        grant_id=submission.grant_id,
        action="portal.form.completed",
        summary=f"{contact_display_name(contact)} completed {template.name}",
    )
    log_event(
        db,
        actor_user_id=None,
        action="portal.form.completed",
        entity_type="portal_form_submission",
        entity_id=str(submission.id),
        meta={
            "case_id": str(submission.case_id),
            "contact_id": str(contact.id),
            "snapshot_file_id": str(snapshot.id),
        },
    )
    notify_portal_staff_form_completed(
        db,
        case_id=submission.case_id,
        contact=contact,
        form_name=template.name,
        matter_label=client_matter_description(case),
    )
    return submission


def submission_out(db: Session, submission: PortalFormSubmission) -> dict[str, Any]:
    template = db.get(PortalFormTemplate, submission.template_id)
    contact = db.get(Contact, submission.contact_id)
    snapshot_name = ""
    if submission.snapshot_file_id:
        f = db.get(File, submission.snapshot_file_id)
        if f:
            snapshot_name = f.original_filename or ""
    return {
        "id": str(submission.id),
        "case_id": str(submission.case_id),
        "template_id": str(submission.template_id),
        "template_name": template.name if template else "",
        "template_reference": template.reference if template else "",
        "contact_id": str(submission.contact_id),
        "contact_name": contact_display_name(contact) if contact else "",
        "status": submission.status.value,
        "responses": submission.responses or {},
        "snapshot_file_id": str(submission.snapshot_file_id) if submission.snapshot_file_id else None,
        "snapshot_filename": snapshot_name,
        "sent_at": submission.sent_at,
        "completed_at": submission.completed_at,
        "voided_at": submission.voided_at,
    }


def portal_form_detail(db: Session, submission: PortalFormSubmission) -> dict[str, Any]:
    template = db.get(PortalFormTemplate, submission.template_id)
    if template is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Template not found")
    fields = load_template_fields(db, template.id)
    base = submission_out(db, submission)
    base["fields"] = [
        {
            "field_key": f.field_key,
            "label": f.label,
            "field_type": f.field_type.value,
            "help_text": f.help_text,
            "required": f.required,
            "sort_order": f.sort_order,
            "select_options": _field_select_options(f),
        }
        for f in fields
    ]
    base["description"] = template.description
    return base


def _scope_summary(db: Session, template: PortalFormTemplate) -> str:
    from app.models import MatterHeadType, MatterSubType

    if template.matter_head_type_id is None and template.matter_sub_type_id is None:
        return "All matter types"
    if template.matter_sub_type_id:
        sub = db.get(MatterSubType, template.matter_sub_type_id)
        head = db.get(MatterHeadType, template.matter_head_type_id) if template.matter_head_type_id else None
        parts = [p for p in [(head.name if head else None), (sub.name if sub else None)] if p]
        return " — ".join(parts) if parts else "Sub-type"
    head = db.get(MatterHeadType, template.matter_head_type_id)
    return f"{head.name} (all sub-types)" if head else "Matter type"


def template_out(db: Session, template: PortalFormTemplate, *, include_fields: bool = False) -> dict[str, Any]:
    count = len(load_template_fields(db, template.id))
    out: dict[str, Any] = {
        "id": template.id,
        "name": template.name,
        "reference": template.reference,
        "description": template.description,
        "matter_head_type_id": template.matter_head_type_id,
        "matter_sub_type_id": template.matter_sub_type_id,
        "scope_summary": _scope_summary(db, template),
        "field_count": count,
        "created_at": template.created_at,
        "updated_at": template.updated_at,
    }
    if include_fields:
        out["fields"] = [
            {
                "id": f.id,
                "field_key": f.field_key,
                "label": f.label,
                "field_type": f.field_type.value,
                "help_text": f.help_text,
                "required": f.required,
                "sort_order": f.sort_order,
                "select_options": _field_select_options(f),
            }
            for f in load_template_fields(db, template.id)
        ]
    return out


def _validate_scope(db: Session, mh: uuid.UUID | None, ms: uuid.UUID | None) -> None:
    from app.fee_scale_service import _validate_scope

    _validate_scope(db, mh, ms)


def _ensure_reference_unique(db: Session, ref: str, *, exclude_id: uuid.UUID | None = None) -> str:
    q = select(PortalFormTemplate).where(PortalFormTemplate.reference == ref.strip())
    if exclude_id:
        q = q.where(PortalFormTemplate.id != exclude_id)
    if db.execute(q.limit(1)).scalar_one_or_none():
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="Reference already in use")
    return ref.strip()


def _apply_fields(db: Session, template_id: uuid.UUID, fields_in: list[dict[str, Any]]) -> None:
    existing = load_template_fields(db, template_id)
    for row in existing:
        db.delete(row)
    db.flush()
    seen_keys: set[str] = set()
    for i, spec in enumerate(fields_in):
        key = _normalize_field_key(str(spec.get("field_key") or ""))
        if key in seen_keys:
            raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail=f"Duplicate field key: {key}")
        seen_keys.add(key)
        ftype = PortalFormFieldType(str(spec.get("field_type") or "text"))
        select_opts = _normalize_select_options(spec.get("select_options")) if ftype == PortalFormFieldType.select else []
        if ftype == PortalFormFieldType.select and not select_opts:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=f"Dropdown field needs at least one option: {str(spec.get('label') or key)}",
            )
        db.add(
            PortalFormTemplateField(
                id=uuid.uuid4(),
                template_id=template_id,
                field_key=key,
                label=str(spec.get("label") or "").strip()[:500],
                field_type=ftype,
                help_text=(str(spec.get("help_text")).strip()[:2000] if spec.get("help_text") else None),
                required=bool(spec.get("required")),
                sort_order=int(spec.get("sort_order") if spec.get("sort_order") is not None else i),
                select_options=select_opts,
            )
        )
    db.flush()


def create_template(db: Session, *, payload: dict[str, Any], owner: User) -> PortalFormTemplate:
    mh = payload.get("matter_head_type_id")
    ms = payload.get("matter_sub_type_id")
    _validate_scope(db, mh, ms)
    ref = _ensure_reference_unique(db, str(payload.get("reference") or ""))
    now = utcnow()
    row = PortalFormTemplate(
        id=uuid.uuid4(),
        name=str(payload.get("name") or "").strip()[:300],
        reference=ref,
        description=(str(payload.get("description")).strip() if payload.get("description") else None),
        matter_head_type_id=mh,
        matter_sub_type_id=ms,
        owner_id=owner.id,
        created_at=now,
        updated_at=now,
    )
    db.add(row)
    db.flush()
    fields = payload.get("fields") or []
    if fields:
        _apply_fields(db, row.id, fields)
    return row


def update_template(db: Session, *, template: PortalFormTemplate, payload: dict[str, Any]) -> PortalFormTemplate:
    if "reference" in payload and payload["reference"] is not None:
        template.reference = _ensure_reference_unique(db, str(payload["reference"]), exclude_id=template.id)
    if "name" in payload and payload["name"] is not None:
        template.name = str(payload["name"]).strip()[:300]
    if "description" in payload:
        template.description = str(payload["description"]).strip() if payload.get("description") else None
    if "matter_head_type_id" in payload or "matter_sub_type_id" in payload:
        mh = payload.get("matter_head_type_id", template.matter_head_type_id)
        ms = payload.get("matter_sub_type_id", template.matter_sub_type_id)
        _validate_scope(db, mh, ms)
        template.matter_head_type_id = mh
        template.matter_sub_type_id = ms
    if payload.get("fields") is not None:
        _apply_fields(db, template.id, payload["fields"])
    template.updated_at = utcnow()
    db.add(template)
    db.flush()
    return template


def delete_template(db: Session, template: PortalFormTemplate) -> None:
    pending = db.execute(
        select(PortalFormSubmission)
        .where(
            PortalFormSubmission.template_id == template.id,
            PortalFormSubmission.status == PortalFormSubmissionStatus.pending,
        )
        .limit(1)
    ).scalar_one_or_none()
    if pending:
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="Cannot delete template with pending submissions")
    db.delete(template)
    db.flush()
