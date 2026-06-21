"""Create minimal Word (.docx) files for case compose flows."""

from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path
from typing import Any, Mapping

# ---------------------------------------------------------------------------
# Precedent merge codes
# ---------------------------------------------------------------------------

def _precedent_code_suffix(slot: int) -> str:
    """Merge key suffix for additional clients 2–4, e.g. [TITLE] -> [TITLE_2]."""
    return f"_{slot}]"


def _merge_key_with_suffix(code: str, slot: int) -> str:
    if not code.startswith("[") or not code.endswith("]"):
        return code
    return code[:-1] + _precedent_code_suffix(slot)


# Name / company codes that are repeated for additional clients 2, 3 & 4 (see build_merge_fields).
_ADDITIONAL_CLIENT_NAME_CODES: tuple[str, ...] = (
    "[TITLE]",
    "[FIRST_NAME]",
    "[FIRST_INITIAL]",
    "[MIDDLE_NAME]",
    "[MIDDLE_INITIAL]",
    "[LAST_NAME]",
    "[LAST_INITIAL]",
    "[COMPANY_NAME]",
    "[TRADING_NAME]",
)

# Lawyer matter contacts are organisation-only; merge codes for the lawyer row are company / trading only.
_LAWYER_ROW_NAME_CODES: tuple[str, ...] = ("[COMPANY_NAME]", "[TRADING_NAME]")


PRECEDENT_CODES: dict[str, str] = {
    # Person
    "[TITLE]": "Title (e.g. Mr / Mrs / Dr)",
    "[FIRST_NAME]": "First name",
    "[FIRST_INITIAL]": "First initial (e.g. J)",
    "[MIDDLE_NAME]": "Middle name",
    "[MIDDLE_INITIAL]": "Middle initial",
    "[LAST_NAME]": "Surname",
    "[LAST_INITIAL]": "Surname initial",
    # Organisation
    "[COMPANY_NAME]": "Registered company name (optional)",
    "[TRADING_NAME]": "Trading name (required for organisations)",
    # Address (shared)
    "[ADDR1]": "Address line 1",
    "[ADDR2]": "Address line 2",
    "[ADDR3]": "Town / city",
    "[ADDR4]": "County",
    "[POSTCODE]": "Postcode",
    # Case / matter
    "[MATTER_DESCRIPTION]": "Matter description",
    "[CASE_REF]": "Case reference number",
    "[DATE]": "Date when the document is generated (DD/MM/YYYY)",
    "[FEE_EARNER]": "Fee earner display name (from the case fee earner)",
    "[FEE_EARNER_JOB_TITLE]": "Fee earner job title (from the case fee earner user)",
    "[FEE_EARNER_INITIALS]": "Fee earner initials (from the case fee earner user)",
    "[CONTACT_REF]": "Contact's reference (as stored in canary)",
    # Firm (Admin → Firm details); narrow scope — precedents / compose merge only for now.
    "[FIRM_TRADING_NAME]": "Firm trading name",
    "[FIRM_REGISTERED_NAME]": "Registered company name (optional)",
    "[FIRM_ADDR1]": "Firm address line 1",
    "[FIRM_ADDR2]": "Firm address line 2",
    "[FIRM_TOWN_CITY]": "Firm town / city",
    "[FIRM_COUNTY]": "Firm county",
    "[FIRM_POSTCODE]": "Firm postcode",
}

for _slot_num, _slot_label in ((2, "2nd"), (3, "3rd"), (4, "4th")):
    for _base_key in _ADDITIONAL_CLIENT_NAME_CODES:
        _suff_key = _merge_key_with_suffix(_base_key, _slot_num)
        PRECEDENT_CODES[_suff_key] = (
            f"{PRECEDENT_CODES[_base_key]} — additional client {_slot_num} "
            f"({_slot_label} 'Client' matter contact on the case, by date added)"
        )

for _li in range(1, 5):
    for _base_key in _LAWYER_ROW_NAME_CODES:
        _inner = _base_key[1:-1]
        _lk = f"[LAWYER_{_li}_{_inner}]"
        PRECEDENT_CODES[_lk] = (
            f"Lawyer {_li}: {PRECEDENT_CODES[_base_key]} "
            "(among 'Lawyers' matter contacts, by date added; lawyers are organisation contacts)"
        )

for _li in range(1, 5):
    for _cj in range(1, 5):
        for _base_key in _ADDITIONAL_CLIENT_NAME_CODES:
            _inner = _base_key[1:-1]
            _lk = f"[LAWYER_{_li}_CLIENT_{_cj}_{_inner}]"
            PRECEDENT_CODES[_lk] = f"Lawyer {_li}'s linked client {_cj}: {PRECEDENT_CODES[_base_key]}"

# Extra fields on each lawyer-linked client (name/company codes are in the loop above).
_LAWYER_LINKED_CLIENT_EXTRA: tuple[tuple[str, str], ...] = (
    ("NAME", "Display name on the contact card"),
    ("TYPE", "person or organisation"),
    ("EMAIL", "Email"),
    ("PHONE", "Phone"),
    ("ADDR1", "Address line 1"),
    ("ADDR2", "Address line 2"),
    ("ADDR3", "Town / city"),
    ("ADDR4", "County"),
    ("POSTCODE", "Postcode"),
    ("COUNTRY", "Country"),
    ("MATTER_REFERENCE", "Matter-specific reference on this case"),
    ("MATTER_CONTACT_TYPE", "Matter contact type label on this case"),
)

for _li in range(1, 5):
    for _cj in range(1, 5):
        for _inner, _lab in _LAWYER_LINKED_CLIENT_EXTRA:
            _lk = f"[LAWYER_{_li}_CLIENT_{_cj}_{_inner}]"
            PRECEDENT_CODES[_lk] = f"Lawyer {_li}'s linked client {_cj}: {_lab} (Case matter contact)"

# Shorthand: same values as [LAWYER_1_CLIENT_cj_*] (first Lawyers matter contact on the case, by date added).
_LAWYER_CONTACT_CLIENT_ALIAS_INNERS: tuple[str, ...] = tuple(
    c[1:-1] for c in _ADDITIONAL_CLIENT_NAME_CODES
) + tuple(x[0] for x in _LAWYER_LINKED_CLIENT_EXTRA)

for _cj in range(1, 5):
    for _inner in _LAWYER_CONTACT_CLIENT_ALIAS_INNERS:
        _lk = f"[LAWYER_CONTACT_CLIENT_{_cj}_{_inner}]"
        PRECEDENT_CODES[_lk] = (
            f"Same as [LAWYER_1_CLIENT_{_cj}_{_inner}]: first 'Lawyers' matter contact’s linked client {_cj} "
            "(by date added among Lawyers contacts)"
        )

# Explicit “selected in compose” contact (letter/document precedent); filled when a contact is chosen in the UI.
_CONTACT_COMPOSE_STATIC: tuple[tuple[str, str], ...] = (
    ("[CONTACT_NAME]", "Display name on the contact card"),
    ("[CONTACT_TYPE]", "person or organisation"),
    ("[CONTACT_EMAIL]", "Email"),
    ("[CONTACT_PHONE]", "Phone"),
    ("[CONTACT_ADDR1]", "Address line 1"),
    ("[CONTACT_ADDR2]", "Address line 2"),
    ("[CONTACT_ADDR3]", "Town / city"),
    ("[CONTACT_ADDR4]", "County"),
    ("[CONTACT_POSTCODE]", "Postcode"),
    ("[CONTACT_COUNTRY]", "Country"),
    (
        "[CONTACT_MATTER_REFERENCE]",
        "Matter-specific reference (case contact snapshot only; empty for a global directory contact)",
    ),
    (
        "[CONTACT_MATTER_CONTACT_TYPE]",
        "Matter contact type label on this case (case contact only; empty for a global directory contact)",
    ),
)
for _ck, _desc in _CONTACT_COMPOSE_STATIC:
    PRECEDENT_CODES[_ck] = (
        f"Selected contact for this compose: {_desc}. Empty if no contact was chosen in the dialogue."
    )

PRECEDENT_CODES["[CONTACT_LETTER_DEAR]"] = (
    "Selected compose contact: opening “Dear …,” using the card display name (empty if none chosen). "
    "Prefer this over a manual Dear line so person vs organisation stays correct."
)
PRECEDENT_CODES["[CONTACT_ORG_LINES]"] = (
    "Selected compose contact: trading name and registered company name with line breaks between non-empty "
    "parts only; empty for persons (avoids blank lines from separate [TRADING_NAME]/[COMPANY_NAME] paragraphs)."
)
PRECEDENT_CODES["[CONTACT_ADDRESS_BLOCK]"] = (
    "Selected compose contact: address lines with breaks between non-empty parts only "
    "(compact substitute for separate [CONTACT_ADDR1]… paragraphs)."
)
PRECEDENT_CODES["[CONTACT_ORG_AND_ADDRESS_BLOCK]"] = (
    "Selected compose contact: [CONTACT_ORG_LINES] and [CONTACT_ADDRESS_BLOCK] in one block — "
    "one line break between org and address sections only when both exist (no spare blank row for persons)."
)

PRECEDENT_CODES["[PRIMARY_CLIENT_LETTER_DEAR]"] = (
    "Primary letter addressee: opening “Dear …,” using the card display name "
    "(same contact as unsuffixed [ADDR1] / [ADDRESS_BLOCK])."
)
PRECEDENT_CODES["[ORG_LINES]"] = (
    "Primary addressee (slot 1 Client): organisation trading + registered lines with breaks; "
    "empty for persons. Prefer one paragraph containing this token instead of separate [TRADING_NAME]/[COMPANY_NAME]."
)
PRECEDENT_CODES["[ADDRESS_BLOCK]"] = (
    "Primary addressee: address lines with breaks between non-empty parts only "
    "(substitute for separate [ADDR1]…[POSTCODE] paragraphs)."
)
PRECEDENT_CODES["[ORG_AND_ADDRESS_BLOCK]"] = (
    "Primary addressee: organisation lines (if any) plus address in one block — inserts a single line break "
    "between org and address only when both exist. Prefer this instead of separate [ORG_LINES] then [ADDRESS_BLOCK] "
    "paragraphs to avoid a blank line for person contacts when [ORG_LINES] is empty."
)

for _slot_num, _slot_label in ((2, "2nd"), (3, "3rd"), (4, "4th")):
    _olk = _merge_key_with_suffix("[ORG_LINES]", _slot_num)
    PRECEDENT_CODES[_olk] = (
        f"Organisation lines for {_slot_label} Client matter contact (merge-all or suffix slot); "
        "empty for persons."
    )
    _abk = _merge_key_with_suffix("[ADDRESS_BLOCK]", _slot_num)
    PRECEDENT_CODES[_abk] = f"Address block for {_slot_label} Client matter contact; omits blank lines."
    _oak = _merge_key_with_suffix("[ORG_AND_ADDRESS_BLOCK]", _slot_num)
    PRECEDENT_CODES[_oak] = (
        f"Combined org + address block for {_slot_label} Client contact (same rules as [ORG_AND_ADDRESS_BLOCK])."
    )

for _cn in range(1, 5):
    PRECEDENT_CODES[f"[CLIENT_{_cn}_LETTER_DEAR]"] = (
        f"Dear line for Client matter contact {_cn} (by client order on the matter). "
        "Use with merge-all or when listing multiple clients."
    )

for _base_key in _ADDITIONAL_CLIENT_NAME_CODES:
    _inner = _base_key[1:-1]
    _cc_key = f"[CONTACT_{_inner}]"
    PRECEDENT_CODES[_cc_key] = (
        f"Selected contact for this compose: same field as [{_inner}] for person or organisation name parts; "
        f"always the contact picked in the dialogue (including when “merge all clients” fills [{_inner}] from another client)."
    )

PRECEDENT_CODES["[QUOTE_PROPERTY_VALUE]"] = "Property value used for banded fee scales (formatted GBP)."
for _qi in range(1, 26):
    _qtag = f"{_qi:02d}"
    PRECEDENT_CODES[f"[QUOTE_{_qtag}_LABEL]"] = f"Quote table row {_qi}: description."
    PRECEDENT_CODES[f"[QUOTE_{_qtag}_AMOUNT]"] = f"Quote table row {_qi}: main (net/inclusive) amount."
    PRECEDENT_CODES[f"[QUOTE_{_qtag}_VAT]"] = f"Quote table row {_qi}: VAT amount (Plus VAT lines)."
PRECEDENT_CODES["[QUOTE_MAIN_TOTAL]"] = "Sum of main-column amounts on item lines."
PRECEDENT_CODES["[QUOTE_VAT_TOTAL]"] = "Sum of VAT-column amounts on item lines."
PRECEDENT_CODES["[QUOTE_GRAND_TOTAL]"] = "Main column total plus VAT column total."
PRECEDENT_CODES["[INVOICE_NUMBER]"] = "Approved invoice number."
PRECEDENT_CODES["[INVOICE_DATE]"] = "Invoice date (approval date, e.g. 13 June 2026)."
PRECEDENT_CODES["[INVOICE_BILL_TO]"] = "Bill-to name on the invoice."
PRECEDENT_CODES["[INVOICE_NET_TOTAL]"] = "Sum of net amounts on invoice lines."
PRECEDENT_CODES["[INVOICE_VAT_TOTAL]"] = "Sum of VAT amounts on invoice lines."
PRECEDENT_CODES["[INVOICE_TOTAL]"] = "Invoice total (net plus VAT)."
for _ii in range(1, 26):
    _itag = f"{_ii:02d}"
    PRECEDENT_CODES[f"[INVOICE_{_itag}_TYPE]"] = f"Invoice table row {_ii}: line type (Fee, Disbursement, VAT)."
    PRECEDENT_CODES[f"[INVOICE_{_itag}_DESCRIPTION]"] = f"Invoice table row {_ii}: description."
    PRECEDENT_CODES[f"[INVOICE_{_itag}_NET]"] = f"Invoice table row {_ii}: net amount."
    PRECEDENT_CODES[f"[INVOICE_{_itag}_VAT]"] = f"Invoice table row {_ii}: VAT amount."
    PRECEDENT_CODES[f"[INVOICE_{_itag}_TOTAL]"] = f"Invoice table row {_ii}: line total (net plus VAT)."
PRECEDENT_CODES["[COMPLETION_DATE]"] = "Completion statement date (e.g. 13 June 2026)."
PRECEDENT_CODES["[COMPLETION_TOTAL_DEBIT]"] = "Sum of debit column amounts."
PRECEDENT_CODES["[COMPLETION_TOTAL_CREDIT]"] = "Sum of credit column amounts."
PRECEDENT_CODES["[COMPLETION_BALANCE_LABEL]"] = "Balance label (due from / due to client)."
PRECEDENT_CODES["[COMPLETION_BALANCE_AMOUNT]"] = "Balance amount (absolute value)."
for _ci in range(1, 51):
    _ctag = f"{_ci:02d}"
    PRECEDENT_CODES[f"[COMPLETION_{_ctag}_DESCRIPTION]"] = f"Completion table row {_ci}: category or item description."
    PRECEDENT_CODES[f"[COMPLETION_{_ctag}_DEBIT]"] = f"Completion table row {_ci}: debit amount."
    PRECEDENT_CODES[f"[COMPLETION_{_ctag}_CREDIT]"] = f"Completion table row {_ci}: credit amount."


def _zip_archive_basename(path: str) -> str:
    return path.replace("\\", "/").rsplit("/", 1)[-1]


def _find_ooxml_content_types_member(names: list[str]) -> str | None:
    """Return the Zip member name for ``[Content_Types].xml`` (case-insensitive basename)."""

    for n in names:
        if _zip_archive_basename(n).lower() == "[content_types].xml":
            return n
    return None


def _find_ooxml_document_xml_member(names: list[str]) -> str | None:
    """Return the Zip member for ``word/document.xml`` (case-insensitive ``word`` / ``document.xml``)."""

    for n in names:
        norm = n.replace("\\", "/")
        parts = norm.split("/")
        if len(parts) >= 2 and parts[-2].lower() == "word" and parts[-1].lower() == "document.xml":
            return n
    return None


def validate_docx_package_bytes(raw: bytes) -> None:
    """Raise ``ValueError`` with a plain-language message if ``raw`` is not a WordprocessingML (.docx) package."""

    import io
    import zipfile

    if not raw:
        raise ValueError("The file is empty.")
    if not raw.startswith(b"PK"):
        raise ValueError(
            "This file does not look like a .docx — real Office documents are ZIP archives whose bytes start with PK. "
            "Common causes: the browser saved an HTML/login/error page with a .docx name, the download failed, "
            "or the file is not Word format. Open it in Word and use Save As → Word Document (.docx), or use the "
            "Canary-generated Universal-letter-precedent.docx from the backend container."
        )
    try:
        zf = zipfile.ZipFile(io.BytesIO(raw))
    except zipfile.BadZipFile as e:
        raise ValueError(
            "This is not a valid ZIP archive — the .docx may be truncated, corrupted, or incomplete."
        ) from e
    try:
        names = zf.namelist()
        ct_name = _find_ooxml_content_types_member(names)
        if ct_name is None:
            raise ValueError(
                "Missing [Content_Types].xml — this is not a valid Office Open XML (.docx) package. "
                "Re-save from Microsoft Word (or export as .docx from Google Docs). "
                "If you renamed another format to .docx, merge will not work."
            )
        doc_name = _find_ooxml_document_xml_member(names)
        if doc_name is None:
            raise ValueError("Missing word/document.xml — not a valid Word .docx.")
        try:
            ct_raw = zf.read(ct_name)
        except KeyError as e:
            raise ValueError(
                "The archive lists [Content_Types].xml but it could not be read — the file may be corrupted."
            ) from e
        if not ct_raw.strip():
            raise ValueError(
                "[Content_Types].xml is empty — this .docx package is invalid. Re-save the document from Word."
            )
        if not ct_raw.lstrip().startswith(b"<"):
            raise ValueError(
                "[Content_Types].xml is not valid XML — this .docx package is broken. Re-save from Word."
            )
    finally:
        zf.close()


def is_invalid_ooxml_merge_exception(exc: BaseException) -> bool:
    """True when ``exc`` usually means bytes are not a loadable Word .docx (client/template issue, HTTP 400)."""

    import zipfile

    if isinstance(exc, (zipfile.BadZipFile, KeyError, OSError)):
        return True
    try:
        from docx.opc.exceptions import PackageNotFoundError

        if isinstance(exc, PackageNotFoundError):
            return True
    except ImportError:
        pass
    lowered = str(exc).lower()
    if "there is no item named" in lowered and "content_types" in lowered:
        return True
    if "bad zipfile" in lowered or "bad magic number for file header" in lowered:
        return True
    # python-docx / lxml load of corrupt OOXML
    if type(exc).__name__ == "XMLSyntaxError":
        return True
    return False


def _s_str(v: object) -> str:
    return (v or "").strip() if isinstance(v, str) else ""


def _initial_letter(v: object) -> str:
    t = _s_str(v)
    return t[0].upper() if t else ""


def _core_name_company_for_contact(contact: Any | None) -> dict[str, str]:
    """Nine merge keys shared by primary and additional-client slots."""

    if not contact:
        return {k: "" for k in _ADDITIONAL_CLIENT_NAME_CODES}

    contact_type = _s_str(getattr(contact, "type", "person"))
    first = _s_str(getattr(contact, "first_name", None))
    middle = _s_str(getattr(contact, "middle_name", None))
    last = _s_str(getattr(contact, "last_name", None))

    company = _s_str(getattr(contact, "company_name", None))
    if not company and contact_type == "organisation":
        company = _s_str(getattr(contact, "name", None))

    return {
        "[TITLE]": _s_str(getattr(contact, "title", None)),
        "[FIRST_NAME]": first,
        "[FIRST_INITIAL]": _initial_letter(first),
        "[MIDDLE_NAME]": middle,
        "[MIDDLE_INITIAL]": _initial_letter(middle),
        "[LAST_NAME]": last,
        "[LAST_INITIAL]": _initial_letter(last),
        "[COMPANY_NAME]": company,
        "[TRADING_NAME]": _s_str(getattr(contact, "trading_name", None)),
    }


def _contact_type_str(contact: Any) -> str:
    t = getattr(contact, "type", None)
    if t is None:
        return ""
    if hasattr(t, "value"):
        return str(t.value)
    return str(t)


def _letter_dear_line(contact: Any | None) -> str:
    if not contact:
        return ""
    name = _s_str(getattr(contact, "name", None))
    return f"Dear {name}," if name else ""


def _org_lines_block(contact: Any | None) -> str:
    """Trading + registered company lines for organisations only; embedded newlines, no trailing blanks."""

    if not contact:
        return ""
    if _contact_type_str(contact) != "organisation":
        return ""
    tr = _s_str(getattr(contact, "trading_name", None))
    reg = _s_str(getattr(contact, "company_name", None))
    return "\n".join(x for x in (tr, reg) if x)


def _address_block_lines(contact: Any | None) -> str:
    """Single-string address with line breaks; skips empty parts (no blank lines)."""

    if not contact:
        return ""
    parts = (
        _s_str(getattr(contact, "address_line1", None)),
        _s_str(getattr(contact, "address_line2", None)),
        _s_str(getattr(contact, "city", None)),
        _s_str(getattr(contact, "county", None)),
        _s_str(getattr(contact, "postcode", None)),
    )
    return "\n".join(p for p in parts if p)


def _org_and_address_block(contact: Any | None) -> str:
    """Organisation lines plus address in one string; no extra gap when org is empty (typical for persons)."""

    org = _org_lines_block(contact)
    addr = _address_block_lines(contact)
    if org and addr:
        return f"{org}\n{addr}"
    return org or addr


def _fill_full_client_composite_slots(out: dict[str, str], oc: list[Any]) -> None:
    """ORG_LINES / ADDRESS_BLOCK / dear lines for each Client slot (merge-all layout)."""

    for i, cc in enumerate(oc[:4]):
        slot = i + 1
        org_b = _org_lines_block(cc)
        addr_b = _address_block_lines(cc)
        dear = _letter_dear_line(cc)
        out[f"[CLIENT_{slot}_LETTER_DEAR]"] = dear
        if slot == 1:
            out["[ORG_LINES]"] = org_b
            out["[ADDRESS_BLOCK]"] = addr_b
            out["[ORG_AND_ADDRESS_BLOCK]"] = _org_and_address_block(cc)
            out["[PRIMARY_CLIENT_LETTER_DEAR]"] = dear
        else:
            out[_merge_key_with_suffix("[ORG_LINES]", slot)] = org_b
            out[_merge_key_with_suffix("[ADDRESS_BLOCK]", slot)] = addr_b
            out[_merge_key_with_suffix("[ORG_AND_ADDRESS_BLOCK]", slot)] = _org_and_address_block(cc)


def _set_primary_addressee_composites(out: dict[str, str], contact: Any | None) -> None:
    """Primary unsuffixed address/org composites from the letter addressee row."""

    if contact is None:
        return
    out["[ORG_LINES]"] = _org_lines_block(contact)
    out["[ADDRESS_BLOCK]"] = _address_block_lines(contact)
    out["[ORG_AND_ADDRESS_BLOCK]"] = _org_and_address_block(contact)
    out["[PRIMARY_CLIENT_LETTER_DEAR]"] = _letter_dear_line(contact)


def _fill_client_dear_lines_and_secondary_composites(out: dict[str, str], oc: list[Any]) -> None:
    """Per-slot Dear lines; suffixed org/address for clients 2–4 (slot 1 primary comes from addressee)."""

    for i, cc in enumerate(oc[:4]):
        slot = i + 1
        out[f"[CLIENT_{slot}_LETTER_DEAR]"] = _letter_dear_line(cc)
        if slot >= 2:
            out[_merge_key_with_suffix("[ORG_LINES]", slot)] = _org_lines_block(cc)
            out[_merge_key_with_suffix("[ADDRESS_BLOCK]", slot)] = _address_block_lines(cc)
            out[_merge_key_with_suffix("[ORG_AND_ADDRESS_BLOCK]", slot)] = _org_and_address_block(cc)


def _lawyer_linked_client_extra_map(contact: Any | None) -> dict[str, str]:
    """Extra merge fields for a lawyer-linked CaseContact (beyond the nine name/company keys)."""

    if not contact:
        return {x[0]: "" for x in _LAWYER_LINKED_CLIENT_EXTRA}
    return {
        "NAME": _s_str(getattr(contact, "name", None)),
        "TYPE": _contact_type_str(contact),
        "EMAIL": _s_str(getattr(contact, "email", None)),
        "PHONE": _s_str(getattr(contact, "phone", None)),
        "ADDR1": _s_str(getattr(contact, "address_line1", None)),
        "ADDR2": _s_str(getattr(contact, "address_line2", None)),
        "ADDR3": _s_str(getattr(contact, "city", None)),
        "ADDR4": _s_str(getattr(contact, "county", None)),
        "POSTCODE": _s_str(getattr(contact, "postcode", None)),
        "COUNTRY": _s_str(getattr(contact, "country", None)),
        "MATTER_REFERENCE": _s_str(getattr(contact, "matter_contact_reference", None)),
        "MATTER_CONTACT_TYPE": _s_str(getattr(contact, "matter_contact_type", None)),
    }


def _empty_precedent_field_map() -> dict[str, str]:
    return {k: "" for k in PRECEDENT_CODES}


def _apply_lawyer_merge_slots(
    out: dict[str, str],
    lawyer_slots: list[tuple[Any, list[Any]] | None] | None,
) -> None:
    if not lawyer_slots:
        return
    for i in range(min(4, len(lawyer_slots))):
        slot = lawyer_slots[i]
        if not slot:
            continue
        law_cc, client_list = slot
        li = i + 1
        law_core = _core_name_company_for_contact(law_cc)
        for merge_key in _LAWYER_ROW_NAME_CODES:
            inner = merge_key[1:-1]
            out[f"[LAWYER_{li}_{inner}]"] = law_core.get(merge_key, "")
        for j, cli in enumerate((client_list or [])[:4]):
            cj = j + 1
            ccore = _core_name_company_for_contact(cli)
            for merge_key, val in ccore.items():
                inner = merge_key[1:-1]
                out[f"[LAWYER_{li}_CLIENT_{cj}_{inner}]"] = val
            extras = _lawyer_linked_client_extra_map(cli)
            for inner, val in extras.items():
                out[f"[LAWYER_{li}_CLIENT_{cj}_{inner}]"] = val
            if li == 1:
                for merge_key, val in ccore.items():
                    inner = merge_key[1:-1]
                    out[f"[LAWYER_CONTACT_CLIENT_{cj}_{inner}]"] = val
                for inner, val in extras.items():
                    out[f"[LAWYER_CONTACT_CLIENT_{cj}_{inner}]"] = val


def _fill_compose_selected_contact_codes(out: dict[str, str], contact: Any | None) -> None:
    """Fill ``[CONTACT_*]`` keys from the contact chosen in the compose dialogue (if any)."""

    if not contact:
        return
    core = _core_name_company_for_contact(contact)
    for merge_key in _ADDITIONAL_CLIENT_NAME_CODES:
        inner = merge_key[1:-1]
        out[f"[CONTACT_{inner}]"] = core.get(merge_key, "")
    out["[CONTACT_NAME]"] = _s_str(getattr(contact, "name", None))
    out["[CONTACT_TYPE]"] = _contact_type_str(contact)
    out["[CONTACT_EMAIL]"] = _s_str(getattr(contact, "email", None))
    out["[CONTACT_PHONE]"] = _s_str(getattr(contact, "phone", None))
    out["[CONTACT_ADDR1]"] = _s_str(getattr(contact, "address_line1", None))
    out["[CONTACT_ADDR2]"] = _s_str(getattr(contact, "address_line2", None))
    out["[CONTACT_ADDR3]"] = _s_str(getattr(contact, "city", None))
    out["[CONTACT_ADDR4]"] = _s_str(getattr(contact, "county", None))
    out["[CONTACT_POSTCODE]"] = _s_str(getattr(contact, "postcode", None))
    out["[CONTACT_COUNTRY]"] = _s_str(getattr(contact, "country", None))
    out["[CONTACT_MATTER_REFERENCE]"] = _s_str(getattr(contact, "matter_contact_reference", None))
    out["[CONTACT_MATTER_CONTACT_TYPE]"] = _s_str(getattr(contact, "matter_contact_type", None))
    out["[CONTACT_LETTER_DEAR]"] = _letter_dear_line(contact)
    out["[CONTACT_ORG_LINES]"] = _org_lines_block(contact)
    out["[CONTACT_ADDRESS_BLOCK]"] = _address_block_lines(contact)
    out["[CONTACT_ORG_AND_ADDRESS_BLOCK]"] = _org_and_address_block(contact)


def build_merge_fields(
    case: Any,
    fee_earner_name: str = "",
    fee_earner_job_title: str = "",
    fee_earner_initials: str = "",
    merge_date: date | None = None,
    *,
    merge_all_clients: bool = False,
    ordered_client_contacts: list[Any] | None = None,
    selected_contact: Any | None = None,
    selected_client_slot: int | None = None,
    lawyer_slots: list[tuple[Any, list[Any]] | None] | None = None,
    compose_selected_contact: Any | None = None,
    firm: Any | None = None,
) -> dict[str, str]:
    """Build precedent code→value dict.

    * **merge_all_clients** — Fill client 1 from ``ordered_client_contacts[0]`` into unsuffixed
      keys and ``[ADDR*]``; clients 2–4 into ``[TITLE_2]`` … ``[TRADING_NAME_4]``.
      ``[CONTACT_REF]`` is taken from the first client row.

    * **Single Client matter contact** (``selected_client_slot`` 1–4) — Fill only that client’s
      name/company keys (slot 1 unsuffixed; slots 2–4 use ``_2`` … ``_4``). Address and
      ``[CONTACT_REF]`` come from ``selected_contact``.

    * **Global contact or non-Client matter contact** (``selected_client_slot`` is None) — Fill
      unsuffixed name and address keys only; suffixed client keys stay empty.

    * **compose_selected_contact** — When set (the contact chosen in the compose UI), fills
      ``[CONTACT_*]`` codes from that row even when ``merge_all_clients`` is True, so templates
      can address the picked contact separately from unsuffixed client merge keys.

    * **Composite tokens** — ``[ORG_LINES]``, ``[ADDRESS_BLOCK]``, ``[ORG_AND_ADDRESS_BLOCK]`` (combined, no blank row
      when org is empty), suffixed ``[ORG_LINES_2]`` … ``[ORG_AND_ADDRESS_BLOCK_4]``, ``[PRIMARY_CLIENT_LETTER_DEAR]``,
      ``[CLIENT_1_LETTER_DEAR]`` … ``[CLIENT_4_LETTER_DEAR]``, and compose-only ``[CONTACT_LETTER_DEAR]``,
      ``[CONTACT_ORG_LINES]``, ``[CONTACT_ADDRESS_BLOCK]``, ``[CONTACT_ORG_AND_ADDRESS_BLOCK]``
      pack multiple lines into one placeholder with internal line breaks so merged letters do not
      retain empty paragraphs when parts are blank.
    """

    out = _empty_precedent_field_map()

    def finalize(m: dict[str, str]) -> dict[str, str]:
        _apply_lawyer_merge_slots(m, lawyer_slots)
        _fill_compose_selected_contact_codes(m, compose_selected_contact)
        return m

    matter_desc = _s_str(getattr(case, "title", None)) if case else ""
    case_ref = _s_str(getattr(case, "case_number", None)) if case else ""
    d = merge_date or date.today()
    date_str = d.strftime("%d/%m/%Y")

    out["[MATTER_DESCRIPTION]"] = matter_desc
    out["[CASE_REF]"] = case_ref
    out["[DATE]"] = date_str
    out["[FEE_EARNER]"] = fee_earner_name
    out["[FEE_EARNER_JOB_TITLE]"] = fee_earner_job_title
    out["[FEE_EARNER_INITIALS]"] = _s_str(fee_earner_initials)

    if firm is not None:
        out["[FIRM_TRADING_NAME]"] = _s_str(getattr(firm, "trading_name", None))
        out["[FIRM_REGISTERED_NAME]"] = _s_str(getattr(firm, "registered_company_name", None))
        out["[FIRM_ADDR1]"] = _s_str(getattr(firm, "addr_line1", None))
        out["[FIRM_ADDR2]"] = _s_str(getattr(firm, "addr_line2", None))
        out["[FIRM_TOWN_CITY]"] = _s_str(getattr(firm, "town_city", None))
        out["[FIRM_COUNTY]"] = _s_str(getattr(firm, "county", None))
        out["[FIRM_POSTCODE]"] = _s_str(getattr(firm, "postcode", None))

    oc = [c for c in (ordered_client_contacts or [])][:4]

    if merge_all_clients:
        for i, cc in enumerate(oc):
            core = _core_name_company_for_contact(cc)
            if i == 0:
                for k, v in core.items():
                    out[k] = v
                out["[ADDR1]"] = _s_str(getattr(cc, "address_line1", None))
                out["[ADDR2]"] = _s_str(getattr(cc, "address_line2", None))
                out["[ADDR3]"] = _s_str(getattr(cc, "city", None))
                out["[ADDR4]"] = _s_str(getattr(cc, "county", None))
                out["[POSTCODE]"] = _s_str(getattr(cc, "postcode", None))
                out["[CONTACT_REF]"] = _s_str(getattr(cc, "matter_contact_reference", None))
            else:
                slot = i + 1
                for k, v in core.items():
                    out[_merge_key_with_suffix(k, slot)] = v
        _fill_full_client_composite_slots(out, oc)
        return finalize(out)

    if selected_contact is None:
        return finalize(out)

    contact_ref = _s_str(getattr(selected_contact, "matter_contact_reference", None))
    out["[ADDR1]"] = _s_str(getattr(selected_contact, "address_line1", None))
    out["[ADDR2]"] = _s_str(getattr(selected_contact, "address_line2", None))
    out["[ADDR3]"] = _s_str(getattr(selected_contact, "city", None))
    out["[ADDR4]"] = _s_str(getattr(selected_contact, "county", None))
    out["[POSTCODE]"] = _s_str(getattr(selected_contact, "postcode", None))

    if selected_client_slot is None or not (1 <= selected_client_slot <= 4):
        core = _core_name_company_for_contact(selected_contact)
        for k, v in core.items():
            out[k] = v
        out["[CONTACT_REF]"] = contact_ref
        _set_primary_addressee_composites(out, selected_contact)
        if oc:
            _fill_client_dear_lines_and_secondary_composites(out, oc)
        else:
            out["[CLIENT_1_LETTER_DEAR]"] = _letter_dear_line(selected_contact)
        return finalize(out)

    idx = selected_client_slot - 1
    cc = oc[idx] if idx < len(oc) else None
    if cc is None:
        out["[CONTACT_REF]"] = contact_ref
        _set_primary_addressee_composites(out, selected_contact)
        if oc:
            _fill_client_dear_lines_and_secondary_composites(out, oc)
        else:
            out["[CLIENT_1_LETTER_DEAR]"] = _letter_dear_line(selected_contact)
        return finalize(out)

    core = _core_name_company_for_contact(cc)
    if selected_client_slot == 1:
        for k, v in core.items():
            out[k] = v
    else:
        for k, v in core.items():
            out[_merge_key_with_suffix(k, selected_client_slot)] = v

    out["[CONTACT_REF]"] = contact_ref
    _set_primary_addressee_composites(out, selected_contact)
    if oc:
        _fill_client_dear_lines_and_secondary_composites(out, oc)
    else:
        out["[CLIENT_1_LETTER_DEAR]"] = _letter_dear_line(selected_contact)
    return finalize(out)


def _replace_in_text(text: str, fields: dict[str, str]) -> str:
    # Longest keys first so a shorter placeholder can never break a longer token (defensive).
    for code in sorted(fields.keys(), key=len, reverse=True):
        text = text.replace(code, fields[code])
    return text


def _normalize_post_merge_whitespace(text: str) -> str:
    """Trim gaps left when earlier client slots merge empty within the same paragraph.

    Templates often interleave ``[TITLE] [FIRST_NAME] … [TITLE_2] …`` with literal spaces.
    When slot 1 is empty, spaces remain before slot 2; strip leading horizontal whitespace
    per line and collapse doubled spaces/tabs inside non-empty lines.
    """
    lines: list[str] = []
    for ln in text.split("\n"):
        if not ln.strip():
            lines.append("")
            continue
        collapsed = re.sub(r"[ \t]{2,}", " ", ln)
        lines.append(collapsed.lstrip())
    return "\n".join(lines)


# Inner token without brackets, e.g. TITLE, LAST_NAME_3 — for slot detection.
_NAME_CODE_INNERS: frozenset[str] = frozenset(c[1:-1] for c in _ADDITIONAL_CLIENT_NAME_CODES)

# ``[CODE]`` or ``[modifiers:CODE]`` where modifiers are one or more of b, i, u.
_MERGE_TOKEN_RE = re.compile(
    r"\[\s*((?:[biu]+)\s*:\s*)?([A-Z0-9_]+)\s*\]",
    re.IGNORECASE,
)


@dataclass(frozen=True)
class _MergeTextSegment:
    text: str
    bold: bool | None = None
    italic: bool | None = None
    underline: bool | None = None


def _parse_modifier_letters(mod: str | None) -> tuple[bool | None, bool | None, bool | None]:
    if not mod:
        return None, None, None
    letters = {c for c in mod.lower() if c in "biu"}
    return (
        True if "b" in letters else None,
        True if "i" in letters else None,
        True if "u" in letters else None,
    )


def _paragraph_has_modifier_tokens(text: str) -> bool:
    return bool(re.search(r"\[\s*(?:[biu]+)\s*:", text, re.IGNORECASE))


def _merge_token_pattern_for_fields(fields: Mapping[str, str]) -> re.Pattern[str]:
    inners = sorted({k[1:-1] for k in fields}, key=len, reverse=True)
    if not inners:
        return re.compile(r"(?!x)")
    inner_alt = "|".join(re.escape(c) for c in inners)
    return re.compile(rf"\[\s*(?:[biu]+\s*:\s*)?(?:{inner_alt})\s*\]", re.IGNORECASE)


def _paragraph_has_merge_tokens(text: str, fields: Mapping[str, str]) -> bool:
    return _merge_token_pattern_for_fields(fields).search(text) is not None


def _replace_merge_tokens_to_segments(text: str, fields: Mapping[str, str]) -> list[_MergeTextSegment]:
    segments: list[_MergeTextSegment] = []
    pos = 0
    for m in _MERGE_TOKEN_RE.finditer(text):
        if m.start() > pos:
            segments.append(_MergeTextSegment(text[pos : m.start()]))
        inner = m.group(2).upper()
        key = f"[{inner}]"
        if key in fields:
            bold, italic, underline = _parse_modifier_letters(m.group(1))
            value = fields[key]
            if value:
                segments.append(_MergeTextSegment(value, bold, italic, underline))
        else:
            segments.append(_MergeTextSegment(m.group(0)))
        pos = m.end()
    if pos < len(text):
        segments.append(_MergeTextSegment(text[pos:]))
    return segments


def _segments_plain_text(segments: list[_MergeTextSegment]) -> str:
    return "".join(s.text for s in segments)


def _trim_leading_segment_whitespace(segments: list[_MergeTextSegment]) -> list[_MergeTextSegment]:
    out: list[_MergeTextSegment] = []
    trimmed = False
    for seg in segments:
        if trimmed or not seg.text.strip():
            out.append(seg)
            continue
        text = seg.text.lstrip()
        out.append(_MergeTextSegment(text, seg.bold, seg.italic, seg.underline))
        trimmed = True
    return out


def _run_element_formatting(r_el: Any) -> tuple[bool | None, bool | None, bool | None]:
    """Read direct bold / italic / underline from a ``w:r`` element."""
    from docx.oxml.ns import qn

    rpr = r_el.find(qn("w:rPr"))
    if rpr is None:
        return None, None, None

    def _tri_state(tag: str) -> bool | None:
        el = rpr.find(qn(tag))
        if el is None:
            return None
        val = el.get(qn("w:val"))
        if val is None or val in ("1", "true", "on"):
            return True
        if val in ("0", "false", "off"):
            return False
        return True

    return _tri_state("w:b"), _tri_state("w:i"), _tri_state("w:u")


def _paragraph_to_formatted_segments(para: Any) -> list[_MergeTextSegment]:
    """Extract paragraph text as formatted segments (includes runs inside ``w:hyperlink``)."""
    from docx.oxml.ns import qn

    segments: list[_MergeTextSegment] = []
    for r in para._p.iter(qn("w:r")):
        parts: list[str] = []
        for child in r:
            tag = child.tag
            if tag == qn("w:t"):
                if child.text:
                    parts.append(child.text)
            elif tag == qn("w:tab"):
                parts.append("\t")
            elif tag in (qn("w:br"), qn("w:cr")):
                parts.append("\n")
            elif tag == qn("w:noBreakHyphen"):
                parts.append("\u2011")
            elif tag == qn("w:softHyphen"):
                parts.append("\u00ad")
        text = "".join(parts)
        if not text:
            continue
        bold, italic, underline = _run_element_formatting(r)
        segments.append(_MergeTextSegment(text, bold, italic, underline))
    return _coalesce_formatted_segments(segments)


def _coalesce_formatted_segments(segments: list[_MergeTextSegment]) -> list[_MergeTextSegment]:
    if not segments:
        return []
    out: list[_MergeTextSegment] = []
    cur = segments[0]
    for seg in segments[1:]:
        if seg.bold == cur.bold and seg.italic == cur.italic and seg.underline == cur.underline:
            cur = _MergeTextSegment(cur.text + seg.text, cur.bold, cur.italic, cur.underline)
        else:
            out.append(cur)
            cur = seg
    out.append(cur)
    return out


def _slice_formatted_segments(
    segments: list[_MergeTextSegment],
    start: int,
    end: int,
) -> list[_MergeTextSegment]:
    if start >= end:
        return []
    result: list[_MergeTextSegment] = []
    pos = 0
    for seg in segments:
        seg_start = pos
        seg_end = pos + len(seg.text)
        pos = seg_end
        if seg_end <= start or seg_start >= end:
            continue
        slice_start = max(start, seg_start) - seg_start
        slice_end = min(end, seg_end) - seg_start
        result.append(
            _MergeTextSegment(seg.text[slice_start:slice_end], seg.bold, seg.italic, seg.underline)
        )
    return _coalesce_formatted_segments(result)


def _formatting_for_segment_range(
    segments: list[_MergeTextSegment],
    start: int,
    end: int,
) -> tuple[bool | None, bool | None, bool | None]:
    sliced = _slice_formatted_segments(segments, start, end)
    if not sliced:
        return None, None, None
    b0, i0, u0 = sliced[0].bold, sliced[0].italic, sliced[0].underline
    for seg in sliced[1:]:
        if seg.bold != b0 or seg.italic != i0 or seg.underline != u0:
            return None, None, None
    return b0, i0, u0


def _replace_merge_tokens_in_formatted_segments(
    segments: list[_MergeTextSegment],
    fields: Mapping[str, str],
) -> list[_MergeTextSegment]:
    """Replace merge tokens while preserving formatting on surrounding static text."""
    text = _segments_plain_text(segments)
    if not _paragraph_has_merge_tokens(text, fields):
        return segments
    result: list[_MergeTextSegment] = []
    pos = 0
    for m in _MERGE_TOKEN_RE.finditer(text):
        if m.start() > pos:
            result.extend(_slice_formatted_segments(segments, pos, m.start()))
        inner = m.group(2).upper()
        key = f"[{inner}]"
        if key in fields:
            bold, italic, underline = _parse_modifier_letters(m.group(1))
            if bold is None and italic is None and underline is None:
                bold, italic, underline = _formatting_for_segment_range(segments, m.start(), m.end())
            value = fields[key]
            if value:
                result.append(_MergeTextSegment(value, bold, italic, underline))
        else:
            result.extend(_slice_formatted_segments(segments, m.start(), m.end()))
        pos = m.end()
    if pos < len(text):
        result.extend(_slice_formatted_segments(segments, pos, len(text)))
    return _coalesce_formatted_segments(result)


def _insert_and_between_adjacent_name_placeholders_in_segments(
    segments: list[_MergeTextSegment],
    sep_flags: dict[tuple[int, int], bool],
) -> list[_MergeTextSegment]:
    """Like :func:`_insert_and_between_adjacent_name_placeholders` but keeps static formatting."""
    text = _segments_plain_text(segments)
    if not sep_flags:
        return segments
    matches = list(_MERGE_TOKEN_RE.finditer(text))
    if len(matches) < 2:
        return segments
    result: list[_MergeTextSegment] = []
    pos = 0
    i = 0
    while i < len(matches):
        m = matches[i]
        result.extend(_slice_formatted_segments(segments, pos, m.start()))
        result.extend(_slice_formatted_segments(segments, m.start(), m.end()))
        pos = m.end()
        if i + 1 < len(matches):
            m2 = matches[i + 1]
            s1 = _name_slot_from_placeholder_inner(m.group(2))
            s2 = _name_slot_from_placeholder_inner(m2.group(2))
            if (
                s1 is not None
                and s2 is not None
                and s2 == s1 + 1
                and sep_flags.get((s1, s2), False)
            ):
                result.append(_MergeTextSegment(" and "))
                pos = m2.start()
        i += 1
    result.extend(_slice_formatted_segments(segments, pos, len(text)))
    return _coalesce_formatted_segments(result)


def _segments_have_direct_formatting(segments: list[_MergeTextSegment]) -> bool:
    return any(
        s.bold is not None or s.italic is not None or s.underline is not None for s in segments
    )


def _name_slot_from_placeholder_inner(inner: str) -> int | None:
    """Return 1–4 for per-client name/company placeholders; ``None`` for other codes."""
    if inner in _NAME_CODE_INNERS:
        return 1
    m = re.fullmatch(r"(.+)_([234])$", inner)
    if not m:
        return None
    base, suf = m.group(1), m.group(2)
    if base not in _NAME_CODE_INNERS:
        return None
    return int(suf)


def _contact_has_any_name_or_company_field(contact: Any | None) -> bool:
    core = _core_name_company_for_contact(contact)
    return any((v or "").strip() for v in core.values())


def _inter_client_sep_flags(ordered_clients: list[Any] | None) -> dict[tuple[int, int], bool]:
    """When True, insert `` and `` between adjacent name placeholders for slots (a, b)."""
    if not ordered_clients:
        return {}
    out: dict[tuple[int, int], bool] = {}
    n = min(len(ordered_clients), 4)
    for i in range(n - 1):
        a, b = ordered_clients[i], ordered_clients[i + 1]
        if _contact_has_any_name_or_company_field(a) and _contact_has_any_name_or_company_field(b):
            out[(i + 1, i + 2)] = True
    return out


def _insert_and_between_adjacent_name_placeholders(
    text: str,
    sep_flags: dict[tuple[int, int], bool],
) -> str:
    """Insert the word ``and`` between consecutive client name placeholders when ``sep_flags`` says to.

    Matches only **adjacent** ``[CODE]`` tokens in this string (same paragraph). Whitespace
    between them is replaced by `` and `` (spaces around *and*).
    """
    if not sep_flags:
        return text
    matches = list(_MERGE_TOKEN_RE.finditer(text))
    if len(matches) < 2:
        return text
    parts: list[str] = []
    pos = 0
    i = 0
    while i < len(matches):
        m = matches[i]
        parts.append(text[pos : m.start()])
        parts.append(m.group(0))
        pos = m.end()
        if i + 1 < len(matches):
            m2 = matches[i + 1]
            s1 = _name_slot_from_placeholder_inner(m.group(2))
            s2 = _name_slot_from_placeholder_inner(m2.group(2))
            if (
                s1 is not None
                and s2 is not None
                and s2 == s1 + 1
                and sep_flags.get((s1, s2), False)
            ):
                parts.append(" and ")
                pos = m2.start()
        i += 1
    parts.append(text[pos:])
    return "".join(parts)


def _xml_escape_ooxml_text(value: str) -> str:
    """Escape text merged into ``<w:t>`` (and similar) XML character data."""
    from xml.sax.saxutils import escape

    return escape(value, {'"': "&quot;", "'": "&apos;"})


def _ooxml_part_paths_for_merge() -> tuple[str, ...]:
    """Part paths inside the .docx zip that may contain visible merge tokens."""
    return (
        "word/document.xml",
        "word/footnotes.xml",
        "word/endnotes.xml",
    )


def _ooxml_part_path_matches(name: str) -> bool:
    if name in _ooxml_part_paths_for_merge():
        return True
    if name.startswith("word/header") and name.endswith(".xml"):
        return True
    if name.startswith("word/footer") and name.endswith(".xml"):
        return True
    return False


def _merge_precedent_codes_in_ooxml_zip(src_bytes: bytes, fields: dict[str, str]) -> bytes:
    """Replace ``[CODE]`` substrings in raw OOXML parts (one pass per file).

    Runs **after** the python-docx paragraph pass. Catches any remaining contiguous
    placeholders in XML (including footnotes) and tokens split across ``<w:t>`` boundaries
    that the paragraph walk could not join for ``and`` insertion.
    """
    import io
    import zipfile

    escaped = {k: _xml_escape_ooxml_text(v) for k, v in fields.items()}
    src = io.BytesIO(src_bytes)
    out = io.BytesIO()
    with zipfile.ZipFile(src, "r") as zin, zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for info in zin.infolist():
            raw = zin.read(info.filename)
            if _ooxml_part_path_matches(info.filename):
                try:
                    text = raw.decode("utf-8")
                except UnicodeDecodeError:
                    zout.writestr(info, raw)
                    continue
                text = _replace_merge_tokens_in_ooxml_text(text, escaped)
                raw = text.encode("utf-8")
            zout.writestr(info, raw)
    return out.getvalue()


def _replace_merge_tokens_in_ooxml_text(text: str, fields: Mapping[str, str]) -> str:
    """Replace merge tokens in raw OOXML text (modifiers are dropped — plain escaped value)."""

    def repl(m: re.Match[str]) -> str:
        key = f"[{m.group(2).upper()}]"
        if key not in fields:
            return m.group(0)
        return fields[key]

    return _MERGE_TOKEN_RE.sub(repl, text)


def _rewrite_paragraph_to_runs(para: Any, segments: list[_MergeTextSegment]) -> None:
    """Replace paragraph content with formatted runs (supports embedded ``\\n`` as line breaks)."""
    from docx.enum.text import WD_BREAK
    from docx.oxml.ns import qn

    def _apply_run_formatting(run: Any, seg: _MergeTextSegment) -> None:
        if seg.bold is not None:
            run.bold = seg.bold
        if seg.italic is not None:
            run.italic = seg.italic
        if seg.underline is not None:
            run.underline = seg.underline

    p_el = para._p
    for child in list(p_el):
        if child.tag != qn("w:pPr"):
            p_el.remove(child)
    for seg in segments:
        if not seg.text:
            continue
        has_fmt = seg.bold is not None or seg.italic is not None or seg.underline is not None
        lines = seg.text.split("\n")
        if has_fmt and len(lines) > 1:
            run = para.add_run(lines[0])
            _apply_run_formatting(run, seg)
            for line in lines[1:]:
                run.add_break(WD_BREAK.LINE)
                if line:
                    run.add_text(line)
            continue
        for i, part in enumerate(lines):
            if i > 0:
                br_run = para.add_run()
                br_run.add_break(WD_BREAK.LINE)
                if has_fmt:
                    _apply_run_formatting(br_run, seg)
            if not part:
                continue
            run = para.add_run(part)
            _apply_run_formatting(run, seg)


def _rewrite_paragraph_to_single_run(para: Any, replaced: str) -> None:
    """Replace paragraph content with plain text (supports embedded ``\\n`` as Word line breaks).

    Word often puts merge tokens in hyperlinked or oddly split runs. Clearing only ``para.runs``
    can leave ``w:hyperlink`` / nested ``w:t`` behind, so the old token still appears next to
    the merged text (e.g. surname twice with a gap). We strip non-``w:pPr`` children and add
    fresh runs — same approach as a clean retype of the paragraph.
    """
    from docx.enum.text import WD_BREAK
    from docx.oxml.ns import qn

    p_el = para._p
    for child in list(p_el):
        if child.tag != qn("w:pPr"):
            p_el.remove(child)
    parts = replaced.split("\n")
    for i, part in enumerate(parts):
        if i > 0:
            para.add_run().add_break(WD_BREAK.LINE)
        para.add_run(part)


def _copy_section_page_geometry(src_section: Any, tgt_section: Any) -> None:
    """Apply letterhead section ``w:pgSz`` / ``w:pgMar`` onto the precedent section.

    Header/footer blocks alone do not define vertical padding: the section's page margins and
    ``header`` / ``footer`` distances (python-docx: ``header_distance``, ``footer_distance``)
    reserve space for header/footer content above and below the body. Without copying these,
    the precedent's geometry wins and letterhead spacing set in Word disappears in compose
    (especially noticeable in ONLYOFFICE).
    """

    for attr in (
        "page_width",
        "page_height",
        "orientation",
        "left_margin",
        "right_margin",
        "top_margin",
        "bottom_margin",
        "header_distance",
        "footer_distance",
        "gutter",
    ):
        try:
            val = getattr(src_section, attr)
        except AttributeError:
            continue
        if val is None:
            continue
        setattr(tgt_section, attr, val)


PRECEDENT_BODY_MARKER = "[PRECEDENT_BODY]"
QUOTE_FEE_TABLE_MARKER = "[QUOTE_FEE_TABLE]"
QUOTE_TABLE_MARKERS = (QUOTE_FEE_TABLE_MARKER, PRECEDENT_BODY_MARKER)
QUOTE_MERGE_SLOT_COUNT = 25
_QUOTE_SLOT_TOKEN_RE = re.compile(r"^\[QUOTE_\d{2}_(?:LABEL|AMOUNT|VAT)\]$", re.IGNORECASE)
INVOICE_MERGE_SLOT_COUNT = 25
_INVOICE_SLOT_TOKEN_RE = re.compile(
    r"^\[INVOICE_\d{2}_(?:TYPE|DESCRIPTION|NET|VAT|TOTAL)\]$",
    re.IGNORECASE,
)
COMPLETION_MERGE_SLOT_COUNT = 50
_COMPLETION_SLOT_TOKEN_RE = re.compile(
    r"^\[COMPLETION_\d{2}_(?:DESCRIPTION|DEBIT|CREDIT)\]$",
    re.IGNORECASE,
)


def insert_xlsx_grid_table_at_marker(doc_bytes: bytes, grid: "XlsxGrid") -> bytes:
    """Replace ``[QUOTE_FEE_TABLE]`` / ``[PRECEDENT_BODY]`` with a Word table built from an xlsx grid."""
    import io
    from copy import deepcopy

    from docx import Document
    from docx.oxml.ns import qn

    from app.xlsx_util import XlsxGrid

    if not isinstance(grid, XlsxGrid):
        raise TypeError("grid must be XlsxGrid")

    doc = Document(io.BytesIO(doc_bytes))
    body = doc.element.body
    p_tag = qn("w:p")
    t_tag = qn("w:t")
    tbl_tag = qn("w:tbl")
    sect_pr_tag = qn("w:sectPr")

    marker_para = None
    for el in list(body):
        if el.tag != p_tag:
            continue
        text = "".join((t.text or "") for t in el.iter(t_tag))
        if any(m in text for m in QUOTE_TABLE_MARKERS):
            marker_para = el
            break

    tmp = Document()
    nrows = len(grid.rows) or 1
    ncols = max((len(r) for r in grid.rows), default=1) or 1
    table = tmp.add_table(rows=nrows, cols=ncols)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for r_i, row in enumerate(grid.rows):
        for c_i in range(ncols):
            val = row[c_i] if c_i < len(row) else ""
            cell = table.rows[r_i].cells[c_i]
            cell.text = val
            if (r_i, c_i) in grid.bold:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
    for r0, c0, r1, c1 in grid.merges:
        try:
            table.rows[r0].cells[c0].merge(table.rows[r1].cells[c1])
        except (IndexError, ValueError):
            pass

    tbl_el = deepcopy(table._tbl)

    if marker_para is not None:
        parent = marker_para.getparent()
        idx = list(parent).index(marker_para)
        parent.insert(idx, tbl_el)
        parent.remove(marker_para)
    else:
        sect_pr = next((el for el in body if el.tag == sect_pr_tag), None)
        if sect_pr is not None:
            idx = list(body).index(sect_pr)
            body.insert(idx, tbl_el)
        else:
            body.append(tbl_el)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def format_gbp_pence(pence: int | None) -> str:
    if pence is None:
        return ""
    negative = pence < 0
    pence = abs(pence)
    pounds = pence / 100
    text = f"£{pounds:,.2f}"
    return f"-{text}" if negative else text


def insert_quote_fee_table_at_marker(
    doc_bytes: bytes,
    rows: list[tuple[str, str | None, bool, bool]],
) -> bytes:
    """Insert a two-column fee table. Each row: (label, amount_display, is_bold, amount_right)."""
    import io
    from copy import deepcopy

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document(io.BytesIO(doc_bytes))
    body = doc.element.body
    p_tag = qn("w:p")
    t_tag = qn("w:t")
    sect_pr_tag = qn("w:sectPr")

    marker_para = None
    for el in list(body):
        if el.tag != p_tag:
            continue
        text = "".join((t.text or "") for t in el.iter(t_tag))
        if any(m in text for m in QUOTE_TABLE_MARKERS):
            marker_para = el
            break

    tmp = Document()
    nrows = max(len(rows), 1)
    table = tmp.add_table(rows=nrows, cols=2)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    data = rows if rows else [(" ", "")]
    for r_i, (label, amount, is_bold, amount_right) in enumerate(data):
        left = table.rows[r_i].cells[0]
        right = table.rows[r_i].cells[1]
        left.text = label
        right.text = amount or ""
        for cell, right_align in ((left, False), (right, amount_right)):
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT if right_align else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    if is_bold:
                        run.bold = True

    tbl_el = deepcopy(table._tbl)
    if marker_para is not None:
        parent = marker_para.getparent()
        idx = list(parent).index(marker_para)
        parent.insert(idx, tbl_el)
        parent.remove(marker_para)
    else:
        sect_pr = next((el for el in body if el.tag == sect_pr_tag), None)
        if sect_pr is not None:
            idx = list(body).index(sect_pr)
            body.insert(idx, tbl_el)
        else:
            body.append(tbl_el)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def strip_precedent_body_marker(doc_bytes: bytes) -> bytes:
    """Remove any leftover ``[PRECEDENT_BODY]`` token from a .docx so it never renders literally.

    Replaces the token text with an empty string inside any paragraph whose combined run-text
    contains it (keeps the surrounding paragraph for layout / spacing). Used as a safety net for
    the blank-letter compose path where the splice is skipped, and as defence-in-depth after the
    splice. Handles the marker even when split across multiple ``w:t`` runs (a common artefact of
    editing in ONLYOFFICE / Word).
    """
    import io
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(io.BytesIO(doc_bytes))
    body = doc.element.body
    p_tag = qn("w:p")
    t_tag = qn("w:t")
    changed = False
    for el in list(body.iter(p_tag)):
        runs_text = "".join((t.text or "") for t in el.iter(t_tag))
        if PRECEDENT_BODY_MARKER not in runs_text:
            continue
        cleaned = runs_text.replace(PRECEDENT_BODY_MARKER, "")
        for r in list(el.findall(qn("w:r"))):
            el.remove(r)
        if cleaned:
            r = el.makeelement(qn("w:r"), {})
            t = el.makeelement(qn("w:t"), {qn("xml:space"): "preserve"})
            t.text = cleaned
            r.append(t)
            el.append(r)
        changed = True
    if not changed:
        return doc_bytes
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def splice_precedent_into_blank_letter(blank_letter_bytes: bytes, precedent_bytes: bytes) -> bytes:
    """Use BLANK_LETTER as the scaffold and inject the chosen precedent's body content into it.

    BLANK_LETTER provides the letter shell: headers, footers, page geometry, and the body merge-code
    scaffold (e.g. recipient address block, date, ``Your Ref`` / ``Our Ref``, salutation, ``Re:``).
    The chosen precedent contributes only its body block elements (paragraphs + tables); its own
    headers/footers/page geometry and any trailing ``w:sectPr`` are discarded.

    Insertion point:
      - If BLANK_LETTER contains a paragraph whose visible text includes ``[PRECEDENT_BODY]``,
        that paragraph is **replaced** with the precedent body elements. This is the recommended
        way to position the precedent body precisely (e.g. between salutation and a static signature
        block that lives in BLANK_LETTER).
      - Otherwise the precedent body is appended at the **end of BLANK_LETTER's body**, just before
        the trailing ``w:sectPr`` page-setup element. For a typical scaffold that ends with
        ``Re: …``, this puts the chosen precedent's content immediately after the subject line.

    Caveats:
      - Style references in the precedent body (e.g. ``Heading 1``) resolve against BLANK_LETTER's
        ``word/styles.xml``. Common built-in styles work; precedent-only custom styles fall back to
        defaults.
      - Numbering definitions and embedded images in the precedent body may not transfer (numbered
        lists can lose their numbering format; image rels may dangle). For anything that must always
        render correctly, put it in BLANK_LETTER.
      - Merge-code substitution must run on the combined result (caller's responsibility).
    """
    import io
    from copy import deepcopy

    from docx import Document

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    p_tag = f"{{{W}}}p"
    tbl_tag = f"{{{W}}}tbl"
    t_tag = f"{{{W}}}t"
    sect_pr_tag = f"{{{W}}}sectPr"

    base = Document(io.BytesIO(blank_letter_bytes))
    src = Document(io.BytesIO(precedent_bytes))

    base_body = base.element.body
    src_body = src.element.body

    src_elements = [deepcopy(el) for el in src_body if el.tag in (p_tag, tbl_tag)]
    if not src_elements:
        out_empty = io.BytesIO()
        base.save(out_empty)
        return out_empty.getvalue()

    marker_para = None
    for el in base_body:
        if el.tag != p_tag:
            continue
        text = "".join((t.text or "") for t in el.iter(t_tag))
        if PRECEDENT_BODY_MARKER in text:
            marker_para = el
            break

    if marker_para is not None:
        parent = marker_para.getparent()
        idx = list(parent).index(marker_para)
        for offset, new_el in enumerate(src_elements):
            parent.insert(idx + offset, new_el)
        parent.remove(marker_para)
    else:
        sect_pr = next((el for el in base_body if el.tag == sect_pr_tag), None)
        if sect_pr is not None:
            idx = list(base_body).index(sect_pr)
            for offset, new_el in enumerate(src_elements):
                base_body.insert(idx + offset, new_el)
        else:
            for new_el in src_elements:
                base_body.append(new_el)

    out = io.BytesIO()
    base.save(out)
    return out.getvalue()


_OD_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_HF_PART_RE = re.compile(r"^word/(header|footer)\d+\.xml$")


def _read_docx_zip_parts(raw: bytes) -> dict[str, bytes]:
    import io
    import zipfile

    parts: dict[str, bytes] = {}
    with zipfile.ZipFile(io.BytesIO(raw), "r") as zf:
        for name in zf.namelist():
            parts[name] = zf.read(name)
    return parts


def _write_docx_zip_parts(parts: dict[str, bytes]) -> bytes:
    import io
    import zipfile

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in parts.items():
            zout.writestr(name, data)
    return buf.getvalue()


def _ooxml_rels_part_path(part_path: str) -> str:
    folder, name = part_path.rsplit("/", 1)
    return f"{folder}/_rels/{name}.rels"


def _parse_ooxml_relationships(rels_bytes: bytes) -> list[tuple[str, str, str]]:
    import xml.etree.ElementTree as ET

    root = ET.fromstring(rels_bytes)
    rels: list[tuple[str, str, str]] = []
    for el in root:
        if el.tag.rsplit("}", 1)[-1] != "Relationship":
            continue
        rels.append((el.get("Id") or "", el.get("Type") or "", el.get("Target") or ""))
    return rels


def _resolve_word_part_path(target: str) -> str:
    path = target.lstrip("/")
    if not path.startswith("word/"):
        path = f"word/{path}"
    return path


def _letterhead_default_hf_rels_paths(lh_parts: dict[str, bytes]) -> tuple[str | None, str | None]:
    """Return ``(header_rels_path, footer_rels_path)`` for the letterhead default header/footer."""
    import xml.etree.ElementTree as ET

    doc_rels_path = "word/_rels/document.xml.rels"
    header_rels_path: str | None = None
    footer_rels_path: str | None = None

    if doc_rels_path in lh_parts:
        rid_to_target = {
            rid: target for rid, _typ, target in _parse_ooxml_relationships(lh_parts[doc_rels_path])
        }
        try:
            root = ET.fromstring(lh_parts["word/document.xml"])
        except (KeyError, ET.ParseError):
            root = None
        if root is not None:
            header_rid: str | None = None
            footer_rid: str | None = None
            for el in root.iter(f"{{{_W_NS}}}headerReference"):
                htype = el.get(f"{{{_W_NS}}}type")
                if htype is None or htype == "default":
                    header_rid = el.get(f"{{{_OD_REL_NS}}}id")
                    break
            for el in root.iter(f"{{{_W_NS}}}footerReference"):
                ftype = el.get(f"{{{_W_NS}}}type")
                if ftype is None or ftype == "default":
                    footer_rid = el.get(f"{{{_OD_REL_NS}}}id")
                    break
            if header_rid and header_rid in rid_to_target:
                rels = _ooxml_rels_part_path(_resolve_word_part_path(rid_to_target[header_rid]))
                if rels in lh_parts:
                    header_rels_path = rels
            if footer_rid and footer_rid in rid_to_target:
                rels = _ooxml_rels_part_path(_resolve_word_part_path(rid_to_target[footer_rid]))
                if rels in lh_parts:
                    footer_rels_path = rels

    if header_rels_path is None and "word/_rels/header1.xml.rels" in lh_parts:
        header_rels_path = "word/_rels/header1.xml.rels"
    if footer_rels_path is None and "word/_rels/footer1.xml.rels" in lh_parts:
        footer_rels_path = "word/_rels/footer1.xml.rels"
    return header_rels_path, footer_rels_path


def _media_paths_from_hf_rels(rels_bytes: bytes | None) -> list[str]:
    if not rels_bytes:
        return []
    paths: list[str] = []
    for _rid, typ, target in _parse_ooxml_relationships(rels_bytes):
        if not target:
            continue
        if typ and "image" in typ.lower():
            paths.append(_resolve_word_part_path(target))
    return paths


def _merge_content_types_for_media(content_types_bytes: bytes, lh_content_types_bytes: bytes) -> bytes:
    """Copy missing ``Default`` entries (e.g. png) from the letterhead package."""
    import xml.etree.ElementTree as ET

    root = ET.fromstring(content_types_bytes)
    lh_root = ET.fromstring(lh_content_types_bytes)
    existing_ext = {
        (el.get("Extension") or "").lower()
        for el in root
        if el.tag.rsplit("}", 1)[-1] == "Default"
    }
    for el in lh_root:
        if el.tag.rsplit("}", 1)[-1] != "Default":
            continue
        ext = (el.get("Extension") or "").lower()
        if ext and ext not in existing_ext:
            root.append(el)
            existing_ext.add(ext)
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _merge_letterhead_package_assets(precedent_bytes: bytes, letterhead_bytes: bytes) -> bytes:
    """Copy letterhead ``word/media`` parts and header/footer ``.rels`` into a composed document."""
    lh_parts = _read_docx_zip_parts(letterhead_bytes)
    prec_parts = _read_docx_zip_parts(precedent_bytes)

    lh_hdr_rels_path, lh_ftr_rels_path = _letterhead_default_hf_rels_paths(lh_parts)
    lh_hdr_rels = lh_parts.get(lh_hdr_rels_path) if lh_hdr_rels_path else None
    lh_ftr_rels = lh_parts.get(lh_ftr_rels_path) if lh_ftr_rels_path else None

    def _copy_media(rels_bytes: bytes | None) -> bytes | None:
        if not rels_bytes:
            return rels_bytes
        updated = rels_bytes
        for media_path in _media_paths_from_hf_rels(rels_bytes):
            if media_path not in lh_parts:
                continue
            dest_path = media_path
            if dest_path in prec_parts and prec_parts[dest_path] != lh_parts[media_path]:
                base = dest_path.rsplit("/", 1)[-1]
                dest_path = f"word/media/lh_{base}"
                n = 0
                while dest_path in prec_parts:
                    n += 1
                    dest_path = f"word/media/lh_{n}_{base}"
                old_target = media_path.removeprefix("word/")
                new_target = dest_path.removeprefix("word/")
                updated = updated.replace(
                    f'Target="{old_target}"'.encode(),
                    f'Target="{new_target}"'.encode(),
                )
            prec_parts[dest_path] = lh_parts[media_path]
        return updated

    lh_hdr_rels = _copy_media(lh_hdr_rels)
    lh_ftr_rels = _copy_media(lh_ftr_rels)

    for part_path in list(prec_parts):
        if not _HF_PART_RE.match(part_path):
            continue
        rels_path = _ooxml_rels_part_path(part_path)
        if part_path.startswith("word/header") and lh_hdr_rels:
            prec_parts[rels_path] = lh_hdr_rels
        elif part_path.startswith("word/footer") and lh_ftr_rels:
            prec_parts[rels_path] = lh_ftr_rels

    if "[Content_Types].xml" in prec_parts and "[Content_Types].xml" in lh_parts:
        prec_parts["[Content_Types].xml"] = _merge_content_types_for_media(
            prec_parts["[Content_Types].xml"],
            lh_parts["[Content_Types].xml"],
        )

    return _write_docx_zip_parts(prec_parts)


def apply_digital_letterhead_headers_footers(precedent_bytes: bytes, letterhead_bytes: bytes) -> bytes:
    """Copy header and footer XML from the letterhead .docx onto every section of the precedent .docx.

    Intended for “typical” letterhead: logos and firm lines live in headers/footers; precedent body
    stays in the document story so page 1 shows letterhead + letter content together. Embedded
    images in the letterhead are copied from ``word/media`` and header/footer relationship parts
    are merged into the composed package.

    Also copies **section page geometry** (margins, header/footer offsets, page size) from the
    letterhead's first section onto every precedent section so padding matches the uploaded template.
    """
    import io
    from copy import deepcopy

    from docx import Document

    lh = Document(io.BytesIO(letterhead_bytes))
    prec = Document(io.BytesIO(precedent_bytes))

    def _copy_hf(src_section: Any, tgt_section: Any, *, header: bool) -> None:
        src_el = (src_section.header if header else src_section.footer)._element
        tgt_el = (tgt_section.header if header else tgt_section.footer)._element
        for child in list(tgt_el):
            tgt_el.remove(child)
        for child in src_el:
            tgt_el.append(deepcopy(child))

    src_sec = lh.sections[0]
    for tgt_sec in prec.sections:
        _copy_section_page_geometry(src_sec, tgt_sec)
        _copy_hf(src_sec, tgt_sec, header=True)
        _copy_hf(src_sec, tgt_sec, header=False)

    out = io.BytesIO()
    prec.save(out)
    return _merge_letterhead_package_assets(out.getvalue(), letterhead_bytes)


def _coalesce_split_merge_tokens_in_docx(doc_bytes: bytes) -> bytes:
    """Join ``w:r`` text when ONLYOFFICE/Word split a merge token across runs inside one paragraph."""
    import io

    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(io.BytesIO(doc_bytes))
    t_tag = qn("w:t")
    r_tag = qn("w:r")
    changed = False

    def _needs_coalesce(p_el: Any) -> bool:
        texts = [t.text or "" for t in p_el.iter(t_tag)]
        combined = "".join(texts)
        if not _MERGE_TOKEN_RE.search(combined):
            return False
        for m in _MERGE_TOKEN_RE.finditer(combined):
            token = m.group(0)
            if not any(token in (t.text or "") for t in p_el.iter(t_tag)):
                return True
        return False

    def _coalesce_p(p_el: Any) -> None:
        nonlocal changed
        if not _needs_coalesce(p_el):
            return
        combined = "".join(t.text or "" for t in p_el.iter(t_tag))
        for r in list(p_el.findall(r_tag)):
            p_el.remove(r)
        if combined:
            r = p_el.makeelement(r_tag, {})
            attrs = {qn("xml:space"): "preserve"} if combined.strip() != combined else {}
            t = p_el.makeelement(t_tag, attrs)
            t.text = combined
            r.append(t)
            p_el.append(r)
        changed = True

    def _walk_paragraphs(container: Any) -> None:
        for p_el in container.iter(qn("w:p")):
            _coalesce_p(p_el)

    _walk_paragraphs(doc.element.body)
    for section in doc.sections:
        for hf in (
            section.header,
            section.footer,
            section.even_page_header,
            section.even_page_footer,
            section.first_page_header,
            section.first_page_footer,
        ):
            _walk_paragraphs(hf._element)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _walk_paragraphs(cell._tc)

    if not changed:
        return doc_bytes
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def merge_precedent_codes(
    src_bytes: bytes,
    fields: dict[str, str],
    *,
    ordered_clients: list[Any] | None = None,
    merge_all_clients: bool = False,
) -> bytes:
    """Replace [CODE] placeholders in a .docx (precedent merge).

    1. **python-docx paragraph pass** — when ``merge_all_clients`` is true, inserts the word
       ``and`` between adjacent client name placeholders for consecutive clients that both
       have name/company data; then substitutes fields; handles merged table cells; removes
       code-only blank paragraphs.

    2. **Zip / OOXML pass** — replaces any remaining contiguous ``[CODE]`` or
       ``[modifiers:CODE]`` substrings in document parts (including split tokens not fixed
       in step 1; formatting modifiers may be lost when a token was split across XML nodes).
    """
    sep_flags = _inter_client_sep_flags(ordered_clients) if merge_all_clients else {}
    prepared = _coalesce_split_merge_tokens_in_docx(src_bytes)
    merged = _merge_precedent_codes_via_python_docx(prepared, fields, sep_flags)
    return _merge_precedent_codes_in_ooxml_zip(merged, fields)


def _merge_precedent_codes_via_python_docx(
    src_bytes: bytes,
    fields: dict[str, str],
    sep_flags: dict[tuple[int, int], bool],
) -> bytes:
    """Paragraph walk: optional *and* insertion, field replace, blank-line cleanup."""
    import io
    from docx import Document

    doc = Document(io.BytesIO(src_bytes))

    seen_wp: set[Any] = set()

    def _merge_para(para: Any) -> bool:
        """Merge codes in para. Returns True if the para should be removed (became blank)."""
        wp = para._p
        if wp in seen_wp:
            return False
        # Read run-level formatting before flattening so static bold/italic survives merge.
        segments = _paragraph_to_formatted_segments(para)
        full = _segments_plain_text(segments)
        if not full:
            return False  # already empty — don't touch
        if sep_flags:
            segments = _insert_and_between_adjacent_name_placeholders_in_segments(segments, sep_flags)
            full = _segments_plain_text(segments)
        if not _paragraph_has_merge_tokens(full, fields):
            return False
        # Claim only once we will rewrite, so empty / no-code paragraphs visited from duplicate
        # merged cells can still be processed on a later distinct visit (should not happen, but safe).
        seen_wp.add(wp)
        had_modifier_tokens = _paragraph_has_modifier_tokens(full)
        segments = _replace_merge_tokens_in_formatted_segments(segments, fields)
        if _segments_have_direct_formatting(segments) or had_modifier_tokens:
            segments = _trim_leading_segment_whitespace(segments)
            replaced = _normalize_post_merge_whitespace(_segments_plain_text(segments))
            _rewrite_paragraph_to_runs(para, segments)
        else:
            replaced = _normalize_post_merge_whitespace(_segments_plain_text(segments))
            _rewrite_paragraph_to_single_run(para, replaced)
        # Remove the paragraph if it's now blank (was code-only, value was empty)
        return not replaced.strip()

    def _process_paras(paras: Any) -> None:
        plist = list(paras) if not isinstance(paras, list) else paras
        to_remove = [p for p in plist if _merge_para(p)]
        for p in to_remove:
            p._element.getparent().remove(p._element)

    def _iter_distinct_cells(table: Any):
        """Each physical cell once (merged cells share one ``w:tc`` but span multiple grid slots)."""
        seen_tc: set[Any] = set()
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                if tc in seen_tc:
                    continue
                seen_tc.add(tc)
                yield cell

    def _process_table(table: Any) -> None:
        for cell in _iter_distinct_cells(table):
            _process_paras(cell.paragraphs)
            nested = getattr(cell, "tables", None)
            if nested:
                for nt in nested:
                    _process_table(nt)

    # Body paragraphs (not inside tables)
    _process_paras(doc.paragraphs)

    for table in doc.tables:
        _process_table(table)

    # Headers / footers
    for section in doc.sections:
        for hf in (section.header, section.footer,
                   section.even_page_header, section.even_page_footer,
                   section.first_page_header, section.first_page_footer):
            if hf.is_linked_to_previous:
                continue
            _process_paras(hf.paragraphs)
            for table in hf.tables:
                _process_table(table)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


_QUOTE_TABLE_HEADER_FILL = "D0DAEA"
_QUOTE_TABLE_SECTION_FILL = "EEF2F8"
_QUOTE_TABLE_TOTAL_FILL = "D0DAEA"
_QUOTE_TABLE_BORDER_COLOR = "B8C4D4"
_QUOTE_TABLE_BORDER_LIGHT = "D8DEE8"


def _docx_set_cell_shading(cell: Any, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _docx_set_cell_margin_dxa(
    cell: Any,
    *,
    top: int = 80,
    bottom: int = 80,
    left: int = 120,
    right: int = 120,
) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def _docx_set_cell_borders(
    cell: Any,
    *,
    top: dict[str, object] | None = None,
    bottom: dict[str, object] | None = None,
    left: dict[str, object] | None = None,
    right: dict[str, object] | None = None,
) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side, spec in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        if spec is None:
            continue
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), str(spec.get("val", "single")))
        el.set(qn("w:sz"), str(spec.get("sz", 4)))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), str(spec.get("color", "auto")))
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def _docx_set_table_width_pct(table: Any, pct: int = 5000) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(pct))
    tbl_w.set(qn("w:type"), "pct")
    tbl_pr.append(tbl_w)


def _docx_style_table_paragraph(
    para: Any,
    *,
    bold: bool = False,
    size_pt: float = 10,
    alignment: Any | None = None,
) -> None:
    from docx.shared import Pt

    if alignment is not None:
        para.alignment = alignment
    text = para.text
    para.clear()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)


def _docx_quote_border(*, light: bool = False, strong: bool = False) -> dict[str, object]:
    if strong:
        return {"val": "single", "sz": 8, "color": _QUOTE_TABLE_BORDER_COLOR}
    color = _QUOTE_TABLE_BORDER_LIGHT if light else _QUOTE_TABLE_BORDER_COLOR
    return {"val": "single", "sz": 4, "color": color}


def _find_quote_fee_table(doc: Any) -> Any | None:
    for table in doc.tables:
        if not table.rows:
            continue
        if len(table.rows[0].cells) < 3:
            continue
        hdr = (table.rows[0].cells[0].text or "").strip().lower()
        if hdr == "description":
            return table
    return None


def _style_quote_fee_table_header_row(table: Any) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    col_widths = (Inches(3.85), Inches(1.35), Inches(1.35))
    hdr = table.rows[0].cells
    labels = ("Description", "Amount", "VAT")
    for i, label in enumerate(labels):
        cell = hdr[i]
        cell.text = label
        para = cell.paragraphs[0]
        _docx_style_table_paragraph(
            para,
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT,
        )
        _docx_set_cell_shading(cell, _QUOTE_TABLE_HEADER_FILL)
        _docx_set_cell_margin_dxa(cell)
        _docx_set_cell_borders(
            cell,
            top=_docx_quote_border(strong=True),
            bottom=_docx_quote_border(strong=True),
            left=_docx_quote_border() if i == 0 else None,
            right=_docx_quote_border() if i == 2 else None,
        )
    for i, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = width


def _style_quote_fee_table_data_row(
    row: Any,
    *,
    line_kind: str = "item",
) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    is_section = line_kind == "section_header"
    is_total_row = line_kind in ("subtotal", "total")

    if is_section:
        for ci, cell in enumerate(row.cells[:3]):
            if not cell.paragraphs:
                cell.add_paragraph()
            para = cell.paragraphs[0]
            if ci == 0:
                _docx_style_table_paragraph(para, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            elif para.text.strip():
                para.clear()
            _docx_set_cell_shading(cell, _QUOTE_TABLE_SECTION_FILL)
            _docx_set_cell_margin_dxa(cell)
            _docx_set_cell_borders(
                cell,
                bottom=_docx_quote_border(light=True),
                left=_docx_quote_border() if ci == 0 else None,
                right=_docx_quote_border() if ci == 2 else None,
            )
        return

    fill = _QUOTE_TABLE_TOTAL_FILL if is_total_row else None
    for ci, cell in enumerate(row.cells[:3]):
        if not cell.paragraphs:
            cell.add_paragraph()
        para = cell.paragraphs[0]
        _docx_style_table_paragraph(
            para,
            bold=is_total_row,
            alignment=WD_ALIGN_PARAGRAPH.RIGHT if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT,
        )
        _docx_set_cell_margin_dxa(cell)
        if fill:
            _docx_set_cell_shading(cell, fill)
        _docx_set_cell_borders(
            cell,
            bottom=_docx_quote_border(light=True),
            left=_docx_quote_border() if ci == 0 else None,
            right=_docx_quote_border() if ci == 2 else None,
        )


def apply_quote_table_presentation(doc_bytes: bytes, lines: list[Any]) -> bytes:
    """Apply header, section, and total styling to the merged quote fee table."""
    import io

    from docx import Document

    doc = Document(io.BytesIO(doc_bytes))
    table = _find_quote_fee_table(doc)
    if table is None:
        return doc_bytes

    _docx_set_table_width_pct(table)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    _style_quote_fee_table_header_row(table)

    for ri, line in enumerate(lines, start=1):
        if ri >= len(table.rows):
            break
        kind = line.line_kind.value if hasattr(line.line_kind, "value") else str(line.line_kind)
        _style_quote_fee_table_data_row(
            table.rows[ri],
            line_kind=kind,
        )

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def strip_empty_quote_table_rows(doc_bytes: bytes) -> bytes:
    """Remove fee-table rows with no merged content (empty cells or leftover slot placeholders)."""
    import io

    from docx import Document

    doc = Document(io.BytesIO(doc_bytes))
    changed = False

    def _row_blank(cells: list) -> bool:
        if len(cells) < 2:
            return all(not (c.text or "").strip() for c in cells)
        label = (cells[0].text or "").strip()
        amount = (cells[1].text or "").strip()
        vat = (cells[2].text or "").strip() if len(cells) > 2 else ""
        if (
            _QUOTE_SLOT_TOKEN_RE.match(label)
            or _QUOTE_SLOT_TOKEN_RE.match(amount)
            or _QUOTE_SLOT_TOKEN_RE.match(vat)
        ):
            return True
        return not label and not amount and not vat

    def _process_table(table: Any) -> None:
        nonlocal changed
        remove_indices: list[int] = []
        for ri, row in enumerate(table.rows):
            if ri == 0:
                continue
            if _row_blank(row.cells):
                remove_indices.append(ri)
        for ri in reversed(remove_indices):
            table._tbl.remove(table.rows[ri]._tr)
            changed = True

    for table in doc.tables:
        _process_table(table)
    for section in doc.sections:
        for hf in (
            section.header,
            section.footer,
            section.even_page_header,
            section.even_page_footer,
            section.first_page_header,
            section.first_page_footer,
        ):
            if hf.is_linked_to_previous:
                continue
            for table in hf.tables:
                _process_table(table)

    if not changed:
        return doc_bytes
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _invoice_line_type_label(line_type: str) -> str:
    if line_type == "fee":
        return "Fee"
    if line_type == "disbursement":
        return "Disbursement"
    if line_type == "vat":
        return "VAT"
    return line_type.replace("_", " ").title()


def invoice_line_merge_fields(
    *,
    invoice_number: str,
    invoice_date: date,
    bill_to_name: str | None,
    lines: list[dict[str, object]],
    total_pence: int,
    max_slots: int = INVOICE_MERGE_SLOT_COUNT,
) -> dict[str, str]:
    """Indexed merge codes for invoice line tables in a firm template."""
    fields: dict[str, str] = {
        "[INVOICE_NUMBER]": invoice_number,
        "[INVOICE_DATE]": invoice_date.strftime("%d %B %Y"),
        "[INVOICE_BILL_TO]": (bill_to_name or "").strip(),
        "[INVOICE_TOTAL]": format_gbp_pence(total_pence),
    }
    net_total = vat_total = 0
    used = 0
    for i, raw in enumerate(lines, start=1):
        if i > max_slots:
            break
        used = i
        tag = f"{i:02d}"
        line_type = str(raw.get("line_type") or "")
        description = str(raw.get("description") or "")
        amount_pence = int(raw.get("amount_pence") or 0)
        tax_pence = int(raw.get("tax_pence") or 0)
        gross = amount_pence + tax_pence
        net_total += amount_pence
        vat_total += tax_pence
        fields[f"[INVOICE_{tag}_TYPE]"] = _invoice_line_type_label(line_type)
        fields[f"[INVOICE_{tag}_DESCRIPTION]"] = description
        fields[f"[INVOICE_{tag}_NET]"] = format_gbp_pence(amount_pence)
        fields[f"[INVOICE_{tag}_VAT]"] = format_gbp_pence(tax_pence)
        fields[f"[INVOICE_{tag}_TOTAL]"] = format_gbp_pence(gross)
    fields["[INVOICE_NET_TOTAL]"] = format_gbp_pence(net_total)
    fields["[INVOICE_VAT_TOTAL]"] = format_gbp_pence(vat_total)
    for i in range(used + 1, max_slots + 1):
        tag = f"{i:02d}"
        fields[f"[INVOICE_{tag}_TYPE]"] = ""
        fields[f"[INVOICE_{tag}_DESCRIPTION]"] = ""
        fields[f"[INVOICE_{tag}_NET]"] = ""
        fields[f"[INVOICE_{tag}_VAT]"] = ""
        fields[f"[INVOICE_{tag}_TOTAL]"] = ""
    return fields


def finance_item_completion_rows(item: Any) -> list[tuple[str, int, int]]:
    """Return completion-table rows as (description, debit_pence, credit_pence)."""
    name = (getattr(item, "name", None) or "").strip()
    direction = getattr(item, "direction", "debit")
    amount_pence = getattr(item, "amount_pence", None)
    vat_pence = getattr(item, "vat_pence", None)

    rows: list[tuple[str, int, int]] = []
    if direction == "credit":
        credit = int(amount_pence) if amount_pence else 0
        if name or credit:
            rows.append((name, 0, credit))
        return rows

    debit_net = int(amount_pence) if amount_pence else 0
    debit_vat = int(vat_pence) if vat_pence else 0
    if debit_net:
        rows.append((name, debit_net, 0))
    if debit_vat:
        vat_label = f"VAT on {name}" if debit_net else name
        rows.append((vat_label, debit_vat, 0))
    elif name and not debit_net and name.strip().lower() != "vat":
        rows.append((name, 0, 0))
    return rows


def completion_line_merge_fields(
    *,
    statement_date: date,
    finance: Any,
    max_slots: int = COMPLETION_MERGE_SLOT_COUNT,
) -> dict[str, str]:
    """Indexed merge codes for completion statement tables in the universal template."""
    fields: dict[str, str] = {
        "[COMPLETION_DATE]": statement_date.strftime("%d %B %Y"),
    }
    table_rows: list[tuple[str, str, str]] = []
    total_dr = total_cr = 0
    categories = getattr(finance, "categories", None) or []
    for cat in categories:
        cat_name = (getattr(cat, "name", None) or str(cat)).strip()
        if cat_name:
            table_rows.append((cat_name.upper(), "", ""))
        items = getattr(cat, "items", None) or []
        for item in items:
            for desc, debit_pence, credit_pence in finance_item_completion_rows(item):
                debit_s = format_gbp_pence(debit_pence) if debit_pence else ""
                credit_s = format_gbp_pence(credit_pence) if credit_pence else ""
                total_dr += debit_pence
                total_cr += credit_pence
                table_rows.append((desc, debit_s, credit_s))

    used = 0
    for i, (desc, debit, credit) in enumerate(table_rows, start=1):
        if i > max_slots:
            break
        used = i
        tag = f"{i:02d}"
        fields[f"[COMPLETION_{tag}_DESCRIPTION]"] = desc
        fields[f"[COMPLETION_{tag}_DEBIT]"] = debit
        fields[f"[COMPLETION_{tag}_CREDIT]"] = credit

    balance = total_cr - total_dr
    if balance > 0:
        balance_label = "BALANCE DUE FROM CLIENT"
    elif balance < 0:
        balance_label = "BALANCE DUE TO CLIENT"
    else:
        balance_label = "BALANCE"

    fields["[COMPLETION_TOTAL_DEBIT]"] = format_gbp_pence(total_dr)
    fields["[COMPLETION_TOTAL_CREDIT]"] = format_gbp_pence(total_cr)
    fields["[COMPLETION_BALANCE_LABEL]"] = balance_label
    fields["[COMPLETION_BALANCE_AMOUNT]"] = format_gbp_pence(abs(balance))

    for i in range(used + 1, max_slots + 1):
        tag = f"{i:02d}"
        fields[f"[COMPLETION_{tag}_DESCRIPTION]"] = ""
        fields[f"[COMPLETION_{tag}_DEBIT]"] = ""
        fields[f"[COMPLETION_{tag}_CREDIT]"] = ""
    return fields


def strip_empty_completion_table_rows(doc_bytes: bytes) -> bytes:
    """Remove completion-table rows with no merged content (empty cells or leftover slot placeholders)."""
    import io

    from docx import Document

    doc = Document(io.BytesIO(doc_bytes))
    changed = False

    def _row_blank(cells: list) -> bool:
        texts = [(c.text or "").strip() for c in cells]
        if any(_COMPLETION_SLOT_TOKEN_RE.match(t) for t in texts):
            return True
        return not any(texts)

    def _process_table(table: Any) -> None:
        nonlocal changed
        remove_indices: list[int] = []
        for ri, row in enumerate(table.rows):
            if ri == 0:
                continue
            if _row_blank(row.cells):
                remove_indices.append(ri)
        for ri in reversed(remove_indices):
            table._tbl.remove(table.rows[ri]._tr)
            changed = True

    for table in doc.tables:
        _process_table(table)
    for section in doc.sections:
        for hf in (
            section.header,
            section.footer,
            section.even_page_header,
            section.even_page_footer,
            section.first_page_header,
            section.first_page_footer,
        ):
            if hf.is_linked_to_previous:
                continue
            for table in hf.tables:
                _process_table(table)

    if not changed:
        return doc_bytes
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def strip_empty_invoice_table_rows(doc_bytes: bytes) -> bytes:
    """Remove invoice-table rows with no merged content (empty cells or leftover slot placeholders)."""
    import io

    from docx import Document

    doc = Document(io.BytesIO(doc_bytes))
    changed = False

    def _row_blank(cells: list) -> bool:
        texts = [(c.text or "").strip() for c in cells]
        if any(
            _INVOICE_SLOT_TOKEN_RE.match(t)
            for t in texts
        ):
            return True
        return not any(texts)

    def _process_table(table: Any) -> None:
        nonlocal changed
        remove_indices: list[int] = []
        for ri, row in enumerate(table.rows):
            if ri == 0:
                continue
            if _row_blank(row.cells):
                remove_indices.append(ri)
        for ri in reversed(remove_indices):
            table._tbl.remove(table.rows[ri]._tr)
            changed = True

    for table in doc.tables:
        _process_table(table)
    for section in doc.sections:
        for hf in (
            section.header,
            section.footer,
            section.even_page_header,
            section.even_page_footer,
            section.first_page_header,
            section.first_page_footer,
        ):
            if hf.is_linked_to_previous:
                continue
            for table in hf.tables:
                _process_table(table)

    if not changed:
        return doc_bytes
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


_W_MAIN_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# Match self-closing or opening ``w:lang`` / ``w:themeFontLang`` elements only.
_LANG_ELEM_RE = re.compile(
    r"<(?:(?:w):)?(?:lang|themeFontLang)\b[^>]*/?>|<(?:(?:w):)?(?:lang|themeFontLang)\b[^>]*>",
    re.IGNORECASE,
)
_LANG_ATTR_RE = re.compile(
    r'(\s(?:(?:w):)?(?:val|eastAsia|bidi)=["\'])([^"\']*)(["\'])',
    re.IGNORECASE,
)
_STYLES_OPEN_RE = re.compile(r"(<(?:(?:w):)?styles\b[^>]*>)", re.IGNORECASE)
_SETTINGS_OPEN_RE = re.compile(r"(<(?:(?:w):)?settings\b[^>]*>)", re.IGNORECASE)
_DOC_DEFAULTS_BLOCK = (
    "<w:docDefaults><w:rPrDefault><w:rPr>"
    '<w:lang w:val="en-GB" w:eastAsia="en-GB" w:bidi="en-GB"/>'
    "</w:rPr></w:rPrDefault></w:docDefaults>"
)
_THEME_FONT_LANG_BLOCK = '<w:themeFontLang w:val="en-GB" w:eastAsia="en-GB" w:bidi="en-GB"/>'


def _coerce_lang_attr_to_en_gb(val: str | None) -> str:
    v = (val or "").strip()
    if not v:
        return "en-GB"
    low = v.replace("_", "-").lower()
    if low in ("en-us", "en", "en-us-x"):
        return "en-GB"
    if low == "en-gb":
        return "en-GB"
    return v


def _patch_lang_element_xml(elem_xml: str) -> str:
    def attr_repl(match: re.Match[str]) -> str:
        return match.group(1) + _coerce_lang_attr_to_en_gb(match.group(2)) + match.group(3)

    return _LANG_ATTR_RE.sub(attr_repl, elem_xml)


def _patch_ooxml_lang_text(text: str, *, part_filename: str = "") -> str:
    """Patch language tags in-place without rewriting the whole OOXML tree.

    ElementTree ``tostring`` on large ``styles.xml`` parts rewrites namespaces and can
    bloat or corrupt complex firm letterheads — ONLYOFFICE then fails to open the file.
    """
    text = _LANG_ELEM_RE.sub(lambda m: _patch_lang_element_xml(m.group(0)), text)
    base = part_filename.rsplit("/", 1)[-1]
    if base == "styles.xml" and not re.search(r"<(?:(?:w):)?docDefaults\b", text, re.IGNORECASE):
        m = _STYLES_OPEN_RE.search(text)
        if m:
            text = text[: m.end()] + _DOC_DEFAULTS_BLOCK + text[m.end() :]
    if base == "settings.xml" and not re.search(r"<(?:(?:w):)?themeFontLang\b", text, re.IGNORECASE):
        m = _SETTINGS_OPEN_RE.search(text)
        if m:
            text = text[: m.end()] + _THEME_FONT_LANG_BLOCK + text[m.end() :]
    return text


def _patch_ooxml_lang_bytes(raw: bytes, *, part_filename: str = "") -> bytes:
    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw
    patched = _patch_ooxml_lang_text(text, part_filename=part_filename)
    if patched == text:
        return raw
    return patched.encode("utf-8")


def ensure_docx_proofing_language_en_gb_bytes(src_bytes: bytes) -> bytes:
    """Force British English as the document language across Word OOXML parts.

    ONLYOFFICE and Word read ``w:docDefaults``, ``w:themeFontLang``, and per-run ``w:lang``.
    Without a full pass, DS often shows “English (United States)” and rewrites saves as en-US.
    """
    import io
    import zipfile

    if not src_bytes.startswith(b"PK"):
        return src_bytes
    try:
        zin = zipfile.ZipFile(io.BytesIO(src_bytes), "r")
    except zipfile.BadZipFile:
        return src_bytes

    patches: dict[str, bytes] = {}
    for info in zin.infolist():
        if not info.filename.startswith("word/") or not info.filename.endswith(".xml"):
            continue
        raw = zin.read(info.filename)
        patched = _patch_ooxml_lang_bytes(raw, part_filename=info.filename)
        if patched != raw:
            patches[info.filename] = patched

    if not patches:
        zin.close()
        return src_bytes

    out_buf = io.BytesIO()
    with zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
        for info in zin.infolist():
            data = patches.get(info.filename, zin.read(info.filename))
            zout.writestr(info, data)
    zin.close()
    return out_buf.getvalue()


def normalize_onlyoffice_persisted_docx_bytes(
    data: bytes,
    *,
    filename: str | None = None,
    mime_type: str | None = None,
) -> bytes:
    """Re-save a .docx exported via ONLYOFFICE ``downloadAs`` so it can be reopened reliably.

    ``downloadAs`` can leave orphan relationship parts and OOXML that triggers ONLYOFFICE
    ``changesError`` on the next open. Round-tripping through python-docx strips those artefacts
    while preserving body content, tables, headers, and embedded media.
    """
    import io

    name = (filename or "").lower()
    mt = (mime_type or "").split(";", 1)[0].strip().lower()
    if not (name.endswith(".docx") or mt == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"):
        return data
    if not data.startswith(b"PK"):
        return data
    try:
        from docx import Document

        doc = Document(io.BytesIO(data))
        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()
    except Exception:
        return data


def finalize_stored_docx_bytes(
    data: bytes,
    *,
    filename: str | None = None,
    mime_type: str | None = None,
) -> bytes:
    """Apply en-GB language normalisation before persisting a .docx from ONLYOFFICE / WebDAV."""
    name = (filename or "").lower()
    mt = (mime_type or "").split(";", 1)[0].strip().lower()
    if name.endswith(".docx") or mt == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return ensure_docx_proofing_language_en_gb_bytes(data)
    return data


def _set_default_proofing_language_en_gb(doc: Any) -> None:
    """Set OOXML default run language to en-GB for new documents.

    ONLYOFFICE/Word use ``w:docDefaults`` for default document / proofing language. python-docx often
    omits ``docDefaults`` until we create it.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    styles_el = doc.styles.element
    dd = styles_el.find(qn("w:docDefaults"))
    if dd is None:
        dd = OxmlElement("w:docDefaults")
        styles_el.insert(0, dd)
    rpd = dd.find(qn("w:rPrDefault"))
    if rpd is None:
        rpd = OxmlElement("w:rPrDefault")
        dd.insert(0, rpd)
    rpr = rpd.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        rpd.append(rpr)
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "en-GB")
    lang.set(qn("w:eastAsia"), "en-GB")
    lang.set(qn("w:bidi"), "en-GB")


def write_blank_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    _set_default_proofing_language_en_gb(doc)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def write_quote_email_precedent_docx(path: Path) -> None:
    """Plain-text-friendly e-mail body for sending a quote (Thunderbird / mailto / Graph)."""
    from docx import Document

    doc = Document()
    _set_default_proofing_language_en_gb(doc)

    lines = [
        "[CONTACT_LETTER_DEAR]",
        "",
        "Thank you for your enquiry. Please find attached our quote for [MATTER_DESCRIPTION].",
        "",
        (
            "The quote sets out the work we propose to undertake and our fees. It also shows VAT and "
            "any disbursements where applicable. If you would like to proceed, or if you have any "
            "questions, please reply to this e-mail."
        ),
        "",
        "We look forward to hearing from you.",
        "",
        "",
        "Kind regards",
        "",
        "[FEE_EARNER]",
        "[FEE_EARNER_JOB_TITLE]",
        "[FIRM_TRADING_NAME]",
        "",
        "---",
        "Our ref: [FEE_EARNER_INITIALS]/[CASE_REF]    Your ref: [CONTACT_REF]",
        "[DATE]",
    ]
    for line in lines:
        doc.add_paragraph(line)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def write_quote_template_docx(path: Path, *, slots: int = QUOTE_MERGE_SLOT_COUNT) -> None:
    """Write a minimal quote .docx with indexed fee-table merge slots (no letter precedent)."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    _set_default_proofing_language_en_gb(doc)

    doc.add_paragraph("[ORG_AND_ADDRESS_BLOCK]")
    doc.add_paragraph("")
    doc.add_paragraph("[DATE]")
    doc.add_paragraph("")
    doc.add_paragraph("Re: [MATTER_DESCRIPTION] — [CASE_REF]")
    doc.add_paragraph("")
    doc.add_paragraph("Dear [CONTACT_LETTER_DEAR]")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Thank you for instructing us. Set out below is our estimate of costs based on a property value of "
        "[QUOTE_PROPERTY_VALUE]."
    )
    doc.add_paragraph("")

    table = doc.add_table(rows=1 + slots, cols=3)
    _docx_set_table_width_pct(table)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    _style_quote_fee_table_header_row(table)

    for i in range(1, slots + 1):
        tag = f"{i:02d}"
        row = table.rows[i].cells
        row[0].text = f"[QUOTE_{tag}_LABEL]"
        row[1].paragraphs[0].text = f"[QUOTE_{tag}_AMOUNT]"
        row[2].paragraphs[0].text = f"[QUOTE_{tag}_VAT]"
        _style_quote_fee_table_data_row(table.rows[i])

    doc.add_paragraph("")
    closing = doc.add_paragraph("Yours faithfully")
    closing.runs[0].font.size = Pt(11)
    doc.add_paragraph("")
    doc.add_paragraph("[FEE_EARNER]")
    doc.add_paragraph("[FIRM_TRADING_NAME]")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def write_quote_template_docx_bytes(*, slots: int = QUOTE_MERGE_SLOT_COUNT) -> bytes:
    import tempfile

    fd, tmp_name = tempfile.mkstemp(suffix=".docx")
    tmp = Path(tmp_name)
    try:
        import os

        os.close(fd)
        write_quote_template_docx(tmp, slots=slots)
        return tmp.read_bytes()
    finally:
        tmp.unlink(missing_ok=True)


def write_completion_statement_docx(
    path: Path,
    *,
    case_number: str,
    client_name: str | None,
    finance: Any,  # FinanceOut (or dict with .categories list)
) -> None:
    """Write a completion statement .docx from case finance data."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _fmt_pence(p: int | None) -> str:
        if p is None:
            return ""
        val = abs(p) / 100
        return f"\u00a3{val:,.2f}"  # £ with thousands separator

    def _set_cell_shading(cell, fill: str) -> None:
        """Apply a background fill colour (hex) to a table cell."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)

    def _set_cell_borders(cell, top=None, bottom=None, left=None, right=None) -> None:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
            if val:
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"), val.get("val", "single"))
                el.set(qn("w:sz"), str(val.get("sz", 4)))
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), val.get("color", "auto"))
                tcBorders.append(el)
        tcPr.append(tcBorders)

    doc = Document()
    _set_default_proofing_language_en_gb(doc)

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── Title ─────────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("COMPLETION STATEMENT")
    run.bold = True
    run.font.size = Pt(16)

    # ── Sub-header: matter reference + date ───────────────────────────────────
    matter_line = case_number
    if client_name:
        matter_line = f"{case_number} — {client_name}"
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(matter_line)
    sub_run.font.size = Pt(11)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Date: {date.today().strftime('%d %B %Y')}")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # spacer

    # ── Main table ────────────────────────────────────────────────────────────
    # Columns: Description | Debit | Credit
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, label in enumerate(("Description", "Debit", "Credit")):
        cell = hdr_cells[i]
        cell.text = label
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
        _set_cell_shading(cell, "D0DAEA")

    # Column widths (Description wide, Debit/Credit equal)
    col_widths = [Inches(3.8), Inches(1.5), Inches(1.5)]
    for i, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = width

    total_dr = 0
    total_cr = 0

    categories = getattr(finance, "categories", None) or []

    for cat in categories:
        cat_name = getattr(cat, "name", None) or str(cat)
        items = getattr(cat, "items", None) or []

        # Category header row
        row = table.add_row()
        row.cells[0].merge(row.cells[2])
        merged = row.cells[0]
        merged.text = cat_name.upper()
        run = merged.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        _set_cell_shading(merged, "EEF2F8")

        for item in items:
            for desc, debit_pence, credit_pence in finance_item_completion_rows(item):
                debit_str = _fmt_pence(debit_pence) if debit_pence else ""
                credit_str = _fmt_pence(credit_pence) if credit_pence else ""
                total_dr += debit_pence
                total_cr += credit_pence

                row = table.add_row()
                row.cells[0].text = desc
                row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
                row.cells[1].text = debit_str
                row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                row.cells[1].paragraphs[0].runs[0 if row.cells[1].paragraphs[0].runs else -1].font.size = Pt(10) if row.cells[1].paragraphs[0].runs else None
                row.cells[2].text = credit_str
                row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if row.cells[2].paragraphs[0].runs:
                    row.cells[2].paragraphs[0].runs[0].font.size = Pt(10)
                for ci in range(3):
                    if row.cells[ci].paragraphs[0].runs:
                        row.cells[ci].paragraphs[0].runs[0].font.size = Pt(10)

    # ── Totals row ────────────────────────────────────────────────────────────
    tot_row = table.add_row()
    tot_row.cells[0].text = "TOTALS"
    tot_row.cells[1].text = _fmt_pence(total_dr)
    tot_row.cells[2].text = _fmt_pence(total_cr)
    for ci, cell in enumerate(tot_row.cells):
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(cell.text)
        run.bold = True
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
        _set_cell_shading(cell, "D0DAEA")

    # ── Balance row ───────────────────────────────────────────────────────────
    balance = total_cr - total_dr
    bal_row = table.add_row()
    bal_row.cells[0].merge(bal_row.cells[1])
    bal_label = bal_row.cells[0]
    bal_label.text = "BALANCE DUE FROM CLIENT" if balance > 0 else "BALANCE DUE TO CLIENT" if balance < 0 else "BALANCE"
    bal_run = bal_label.paragraphs[0].runs[0] if bal_label.paragraphs[0].runs else bal_label.paragraphs[0].add_run(bal_label.text)
    bal_run.bold = True
    bal_run.font.size = Pt(10)
    _set_cell_shading(bal_label, "EEF2F8")

    bal_val_cell = bal_row.cells[2]
    bal_val_cell.text = _fmt_pence(abs(balance))
    bal_val_run = bal_val_cell.paragraphs[0].runs[0] if bal_val_cell.paragraphs[0].runs else bal_val_cell.paragraphs[0].add_run(bal_val_cell.text)
    bal_val_run.bold = True
    bal_val_run.font.size = Pt(10)
    bal_val_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_cell_shading(bal_val_cell, "EEF2F8")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def extract_plain_text_from_docx_bytes(data: bytes) -> str:
    """Best-effort plain text from a .docx for e-mail body (M365 Graph)."""
    from io import BytesIO

    from docx import Document

    doc = Document(BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = (p.text or "").strip()
                    if t:
                        parts.append(t)
    return "\n\n".join(parts) if parts else ""


def write_client_account_reconcile_report_docx(
    path: Path,
    *,
    firm_trading_name: str,
    firm_registered_name: str | None,
    client_bank_account_name: str | None,
    client_bank_sort_code: str | None,
    client_bank_account_number_last4: str | None,
    period_end_date: date,
    ledger_client_total_pence: int,
    ledger_office_total_pence: int,
    bank_statement_balance_pence: int,
    difference_pence: int,
    prepared_by_name: str | None,
    prepared_at: datetime | None,
    approved_by_name: str | None,
    approved_at: datetime | None,
    notes: str | None,
    status: str,
) -> None:
    """Write a client account reconcile report .docx for month-end sign-off."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    def _fmt_pence(p: int) -> str:
        val = p / 100
        sign = "-" if val < 0 else ""
        return f"{sign}£{abs(val):,.2f}"

    def _fmt_dt(dt: datetime | None) -> str:
        if dt is None:
            return "—"
        local = dt
        if local.tzinfo is not None:
            local = local.replace(tzinfo=None)
        return local.strftime("%d %B %Y %H:%M")

    doc = Document()
    _set_default_proofing_language_en_gb(doc)

    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("CLIENT ACCOUNT RECONCILE REPORT")
    run.bold = True
    run.font.size = Pt(16)

    firm_line = (firm_trading_name or "").strip()
    if firm_registered_name and firm_registered_name.strip() and firm_registered_name.strip() != firm_line:
        firm_line = f"{firm_line} ({firm_registered_name.strip()})" if firm_line else firm_registered_name.strip()
    if firm_line:
        firm_para = doc.add_paragraph()
        firm_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        firm_run = firm_para.add_run(firm_line)
        firm_run.font.size = Pt(12)

    period_para = doc.add_paragraph()
    period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    period_run = period_para.add_run(f"Period ended: {period_end_date.strftime('%d %B %Y')}")
    period_run.font.size = Pt(11)

    doc.add_paragraph()

    bank_bits: list[str] = []
    if client_bank_account_name:
        bank_bits.append(client_bank_account_name.strip())
    if client_bank_sort_code:
        bank_bits.append(f"Sort code {client_bank_sort_code.strip()}")
    if client_bank_account_number_last4:
        bank_bits.append(f"Account •••• {client_bank_account_number_last4.strip()}")
    if bank_bits:
        bank_para = doc.add_paragraph("Client bank account: " + " · ".join(bank_bits))
        bank_para.runs[0].font.size = Pt(10)

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    rows_data = [
        ("Ledger client total (all matters)", _fmt_pence(ledger_client_total_pence)),
        ("Bank statement closing balance", _fmt_pence(bank_statement_balance_pence)),
        ("Difference (bank minus ledger)", _fmt_pence(difference_pence)),
        ("Office ledger total (reference)", _fmt_pence(ledger_office_total_pence)),
        ("Status", status.capitalize()),
    ]
    for i, (label, value) in enumerate(rows_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    doc.add_paragraph()

    prep_para = doc.add_paragraph(f"Prepared by: {prepared_by_name or '—'}")
    prep_para.runs[0].font.size = Pt(10)
    prep_at = doc.add_paragraph(f"Prepared at: {_fmt_dt(prepared_at)}")
    prep_at.runs[0].font.size = Pt(10)

    appr_para = doc.add_paragraph(f"Approved by: {approved_by_name or '—'}")
    appr_para.runs[0].font.size = Pt(10)
    appr_at = doc.add_paragraph(f"Approved at: {_fmt_dt(approved_at)}")
    appr_at.runs[0].font.size = Pt(10)

    if notes and notes.strip():
        doc.add_paragraph()
        notes_heading = doc.add_paragraph("Notes")
        notes_heading.runs[0].bold = True
        for line in notes.strip().splitlines():
            doc.add_paragraph(line)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def write_invoice_docx(
    path: Path,
    *,
    firm_trading_name: str,
    firm_registered_name: str | None,
    firm_addr_line1: str | None,
    firm_addr_line2: str | None,
    firm_town_city: str | None,
    firm_county: str | None,
    firm_postcode: str | None,
    invoice_number: str,
    invoice_date: date,
    case_number: str,
    client_name: str | None,
    matter_description: str,
    fee_earner_name: str | None,
    bill_to_name: str | None,
    lines: list[dict[str, object]],
    total_pence: int,
) -> None:
    """Write a client invoice .docx from structured invoice data."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    def _fmt_pence(p: int) -> str:
        val = p / 100
        sign = "-" if val < 0 else ""
        return f"{sign}£{abs(val):,.2f}"

    def _line_type_label(line_type: str) -> str:
        if line_type == "fee":
            return "Fee"
        if line_type == "disbursement":
            return "Disbursement"
        if line_type == "vat":
            return "VAT"
        return line_type.replace("_", " ").title()

    doc = Document()
    _set_default_proofing_language_en_gb(doc)

    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    firm_line = (firm_trading_name or "").strip()
    if firm_registered_name and firm_registered_name.strip() and firm_registered_name.strip() != firm_line:
        firm_line = f"{firm_line} ({firm_registered_name.strip()})" if firm_line else firm_registered_name.strip()
    if firm_line:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(firm_line)
        r.bold = True
        r.font.size = Pt(14)

    addr_parts = [firm_addr_line1, firm_addr_line2, firm_town_city, firm_county, firm_postcode]
    addr_text = ", ".join(x.strip() for x in addr_parts if x and str(x).strip())
    if addr_text:
        addr_para = doc.add_paragraph()
        addr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        addr_run = addr_para.add_run(addr_text)
        addr_run.font.size = Pt(10)
        addr_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("INVOICE")
    title_run.bold = True
    title_run.font.size = Pt(16)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"{invoice_number}  ·  {invoice_date.strftime('%d %B %Y')}"
    )
    meta_run.font.size = Pt(11)

    doc.add_paragraph()

    if bill_to_name and bill_to_name.strip():
        bill_para = doc.add_paragraph()
        bill_para.add_run("Bill to: ").bold = True
        bill_para.add_run(bill_to_name.strip())

    matter_bits = [case_number]
    if client_name and client_name.strip():
        matter_bits.append(client_name.strip())
    if matter_description and matter_description.strip():
        matter_bits.append(matter_description.strip())
    matter_para = doc.add_paragraph("Matter: " + " — ".join(matter_bits))
    matter_para.runs[0].font.size = Pt(10)

    if fee_earner_name and fee_earner_name.strip():
        fe_para = doc.add_paragraph(f"Fee earner: {fee_earner_name.strip()}")
        fe_para.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ("Type", "Description", "Net", "VAT", "Total")
    for i, label in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = label
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 1 else WD_ALIGN_PARAGRAPH.LEFT

    net_total = vat_total = 0
    for raw in lines:
        line_type = str(raw.get("line_type") or "")
        description = str(raw.get("description") or "")
        amount_pence = int(raw.get("amount_pence") or 0)
        tax_pence = int(raw.get("tax_pence") or 0)
        gross = amount_pence + tax_pence
        net_total += amount_pence
        vat_total += tax_pence
        row = table.add_row()
        row.cells[0].text = _line_type_label(line_type)
        row.cells[1].text = description
        row.cells[2].text = _fmt_pence(amount_pence)
        row.cells[3].text = _fmt_pence(tax_pence)
        row.cells[4].text = _fmt_pence(gross)
        for j in range(2, 5):
            row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    summary = doc.add_table(rows=3, cols=2)
    summary.style = "Table Grid"
    summary_rows = [
        ("Net total", _fmt_pence(net_total)),
        ("VAT total", _fmt_pence(vat_total)),
        ("Invoice total", _fmt_pence(total_pence)),
    ]
    for i, (label, value) in enumerate(summary_rows):
        summary.rows[i].cells[0].text = label
        summary.rows[i].cells[1].text = value
        if i == 2:
            summary.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            summary.rows[i].cells[1].paragraphs[0].runs[0].bold = True
        summary.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
